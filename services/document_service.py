from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tkinter import filedialog, simpledialog, messagebox
from models.translations import get_translation

class DocumentService:
    def __init__(self, language="English"):
        self.language = language

    def export_to_docx(self, sentences):
        if not sentences:
            messagebox.showwarning(
                get_translation(self.language, "warning_title"),
                get_translation(self.language, "no_sentences_warning")
            )
            return
        
        title = simpledialog.askstring(
            get_translation(self.language, "document_title_prompt"),
            get_translation(self.language, "document_title_prompt"),
            initialvalue="Fill in the Blank"
        )
        if not title:
            return
        
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
            title=get_translation(self.language, "save_document_title"),
            initialfile=f"{title}.docx"
        )
        
        if not file_path:
            return
        
        try:
            doc = Document()
            
            # Add title
            title_para = doc.add_paragraph(title)
            title_format = title_para.runs[0].font
            title_format.size = Pt(16)
            title_format.bold = True
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add sentences
            doc.add_paragraph()
            for i, (masked, original) in enumerate(sentences, 1):
                p = doc.add_paragraph()
                p.add_run(f"{i}. ").bold = True
                p.add_run(masked)
            
            # Add answer key
            doc.add_page_break()
            key_title = doc.add_paragraph("Answer Key")
            key_title_format = key_title.runs[0].font
            key_title_format.size = Pt(16)
            key_title_format.bold = True
            key_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()
            for i, (masked, original) in enumerate(sentences, 1):
                # Find masked words
                masked_words = set()
                for orig, masked_word in zip(original.split(), masked.split()):
                    if orig != masked_word and masked_word.startswith(orig[0]) and '_' in masked_word:
                        masked_words.add(orig)
                
                p = doc.add_paragraph()
                p.add_run(f"{i}. ").bold = True
                p.add_run(", ".join(sorted(masked_words)))
            
            doc.save(file_path)
            messagebox.showinfo(
                get_translation(self.language, "export_success_title"),
                get_translation(self.language, "export_success_msg")
            )
        except Exception as e:
            messagebox.showerror(
                get_translation(self.language, "error_title"),
                str(e)
            )