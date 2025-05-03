import tkinter as tk
from tkinter import ttk
from models.translations import get_translation, TRANSLATIONS
from tkinter import filedialog, simpledialog, messagebox
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import platform
from nltk import word_tokenize
from models.config import DEFAULT_CONFIG
import requests
from tkinter import scrolledtext
import yaml

# Add these keys to both English and Chinese translations
for language in TRANSLATIONS:
    if "regenerate_button" not in TRANSLATIONS[language]:
        TRANSLATIONS[language]["regenerate_button"] = "↻"
    if "menu_button" not in TRANSLATIONS[language]:
        TRANSLATIONS[language]["menu_button"] = "⋮" 
    if "checkmark" not in TRANSLATIONS[language]:
        TRANSLATIONS[language]["checkmark"] = "✓"
    if "progress_indicator" not in TRANSLATIONS[language]:
        TRANSLATIONS[language]["progress_indicator"] = "{0}/{1}"

class SentenceWidgetManager(ttk.LabelFrame):
    def __init__(self, parent, language, word_processor, api_service, on_sentences_changed=None, main_window=None):
        super().__init__(parent, text=get_translation(language, "generated_sentences"), padding="5")
        self.language = language
        self.word_processor = word_processor
        self.api_service = api_service
        self.sentence_widgets = []
        self.parent = parent
        self.on_sentences_changed = on_sentences_changed
        self.main_window = main_window
        
        # Buttons Frame
        self.buttons_frame = ttk.Frame(self)
        self.buttons_frame.grid(row=0, column=0, sticky=tk.E, padx=5, pady=5)

        # Main menu button (replaced the export button)
        self.menu_btn = ttk.Button(self.buttons_frame, text=get_translation(language, "menu_button_main"), 
                                   command=self._show_main_menu, state="normal")
        self.menu_btn.pack(side=tk.LEFT, padx=(0, 5))
        
        self.show_all_btn = ttk.Button(self.buttons_frame, text=get_translation(language, "show_all"), 
                                     command=self.show_all_words, state="disabled")
        self.show_all_btn.pack(side=tk.LEFT, padx=(0, 5))
        
        self.delete_btn = ttk.Button(self.buttons_frame, text=get_translation(language, "delete_all"), 
                                   command=self.clear_sentences, state="disabled")
        self.delete_btn.pack(side=tk.LEFT)
        
        # Canvas and scrollbar
        self.canvas = tk.Canvas(self)
        self.scrollbar = ttk.Scrollbar(self, orient="vertical", command=self.canvas.yview)
        self.canvas.configure(yscrollcommand=self.scrollbar.set)
        
        self.sentences_container = ttk.Frame(self.canvas)
        self.canvas_frame = self.canvas.create_window((0, 0), window=self.sentences_container, anchor="nw")
        
        self.canvas.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(5, 0))
        self.scrollbar.grid(row=1, column=1, sticky=(tk.N, tk.S))
        
        self.columnconfigure(0, weight=1)
        self.rowconfigure(1, weight=1)
        
        self.sentences_container.bind('<Configure>', self._on_frame_configure)
        self.canvas.bind('<Configure>', self._on_canvas_configure)
        
        # Platform specific mousewheel bindings
        self._bind_mouse_wheel()
        
    def _bind_mouse_wheel(self):
        """Bind mouse wheel scrolling based on platform."""
        system = platform.system()
        
        if system == "Windows":
            # Windows uses <MouseWheel>
            self.canvas.bind("<MouseWheel>", self._on_mousewheel_windows)
            self.bind("<MouseWheel>", self._on_mousewheel_windows)
        elif system == "Darwin":
            # macOS uses <MouseWheel> with different delta values
            self.canvas.bind("<MouseWheel>", self._on_mousewheel_macos)
            self.bind("<MouseWheel>", self._on_mousewheel_macos)
        else:
            # Linux uses Button-4 and Button-5
            self.canvas.bind("<Button-4>", self._on_mousewheel_linux)
            self.canvas.bind("<Button-5>", self._on_mousewheel_linux)
            self.bind("<Button-4>", self._on_mousewheel_linux)
            self.bind("<Button-5>", self._on_mousewheel_linux)
    
    def _on_mousewheel_windows(self, event):
        """Handle mouse wheel scrolling for Windows."""
        if self.canvas.winfo_height() < self.sentences_container.winfo_height():
            self.canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        return "break"
    
    def _on_mousewheel_macos(self, event):
        """Handle mouse wheel scrolling for macOS."""
        if self.canvas.winfo_height() < self.sentences_container.winfo_height():
            self.canvas.yview_scroll(int(-1 * event.delta), "units")
        return "break"
    
    def _on_mousewheel_linux(self, event):
        """Handle mouse wheel scrolling for Linux."""
        if self.canvas.winfo_height() < self.sentences_container.winfo_height():
            if event.num == 4:
                self.canvas.yview_scroll(-1, "units")
            elif event.num == 5:
                self.canvas.yview_scroll(1, "units")
        return "break"
    
    def _on_frame_configure(self, event=None):
        """Update the scrollregion to encompass the inner frame."""
        self.canvas.configure(scrollregion=self.canvas.bbox("all"))
    
    def _on_canvas_configure(self, event):
        """Resize the inner frame to match the canvas."""
        self.canvas.itemconfig(self.canvas_frame, width=event.width)
    
    def _on_mousewheel(self, event):
        """Legacy method, replaced by platform-specific methods."""
        pass
    
    def add_sentence(self, word, sentence):
        frame = ttk.Frame(self.sentences_container)
        frame.grid(sticky=(tk.W, tk.E), pady=2)
        
        # Store original word and sentence for later reference
        frame.original_word = word
        frame.original_sentence = sentence
        
        # Create a 3-column grid layout
        frame.columnconfigure(0, weight=0)  # Order number column fixed width
        frame.columnconfigure(1, weight=1)  # Text column stretches
        frame.columnconfigure(2, weight=0)  # Button column fixed width
        
        # Add order number
        order_number = len(self.sentence_widgets) + 1
        order_label = ttk.Label(frame, text=f"{order_number}.", width=4)
        order_label.grid(row=0, column=0, sticky=tk.W, padx=(5, 0))
        
        # Save order_label reference for updates
        frame.order_label = order_label
        
        # Create masked sentence
        masked_sentence = self._create_masked_sentence(word, sentence)
        
        # Text widget in second column
        text_widget = tk.Text(frame, wrap=tk.WORD, cursor="arrow", height=1, width=50)
        text_widget.insert("1.0", masked_sentence)
        text_widget.configure(state="disabled", selectbackground=text_widget.cget("background"), 
                            selectforeground=text_widget.cget("foreground"), 
                            inactiveselectbackground=text_widget.cget("background"))
        text_widget.grid(row=0, column=1, sticky=(tk.W, tk.E), padx=(0, 5))
        
        # Bind to prevent text selection
        for seq in ["<Button-1>", "<B1-Motion>", "<Double-Button-1>", "<Triple-Button-1>"]:
            text_widget.bind(seq, lambda e: "break")
        
        # Adjust height
        def adjust_text_height():
            text_widget.configure(state="normal")
            text_widget.see("end")
            num_lines = text_widget.count("1.0", "end", "displaylines")[0]
            text_widget.configure(height=max(2, num_lines))
            text_widget.configure(state="disabled")
        
        text_widget.after(10, adjust_text_height)
        text_widget.bind('<Configure>', lambda e: text_widget.after(10, adjust_text_height))
        
        # Create a frame for the buttons in the third column
        buttons_frame = ttk.Frame(frame)
        buttons_frame.grid(row=0, column=2, sticky=tk.E)
        
        # Buttons inside the buttons_frame
        show_btn = ttk.Button(
            buttons_frame,
            text=get_translation(self.language, "show"),
            command=lambda: self._toggle_word(text_widget, show_btn),
            name="show_button"
        )
        show_btn.pack(side=tk.LEFT, padx=(0, 2))
        
        copy_btn = ttk.Button(
            buttons_frame,
            text=get_translation(self.language, "copy"),
            command=lambda: self._copy_sentence(text_widget),
            name="copy_button"
        )
        copy_btn.pack(side=tk.LEFT, padx=(0, 2))
        
        # Store original word for regeneration
        original_word = word
        
        regen_btn = ttk.Button(
            buttons_frame,
            text=get_translation(self.language, "regenerate_button"),
            width=2,
            command=lambda w=original_word, f=frame: self._regenerate_sentence(w, f),
            name="regen_button"
        )
        regen_btn.pack(side=tk.LEFT, padx=(0, 2))
        
        # Create menu button
        menu_btn = ttk.Button(
            buttons_frame,
            text=get_translation(self.language, "menu_button"),
            width=2,
            command=lambda f=frame: self._show_menu(f),
            name="menu_button"
        )
        menu_btn.pack(side=tk.LEFT)
        
        # Store reference to menu button
        frame.menu_btn = menu_btn
        
        # Add the new frame to the list of sentence widgets
        self.sentence_widgets.append(frame)
        
        # Update buttons state
        self._update_buttons_state()
        
        return frame
    
    def _create_masked_sentence(self, word, sentence):
        """Create a masked sentence by identifying and masking the target word."""
        # Get the base form of the input word
        target_word = word.lower()
        
        # Tokenize the sentence
        words = word_tokenize(sentence)
        masked_words = {}
        
        # Process each word in the sentence
        for original_word in words:
            # Skip non-alphabetic words
            if not original_word.isalpha():
                continue
            
            # Check if the word matches the target word or its variations
            if self.word_processor.is_word_match(original_word, target_word):
                masked_word = original_word[0] + '_' * (len(original_word) - 1)
                masked_words[original_word] = masked_word
        
        # Create the masked sentence
        masked_sentence = sentence
        for original, masked in sorted(masked_words.items(), key=lambda x: len(x[0]), reverse=True):
            masked_sentence = masked_sentence.replace(original, masked)
        
        return masked_sentence
    
    def _toggle_word(self, text_widget, button):
        # Find the parent frame
        frame = text_widget.master
        
        # Get the original sentence from the frame
        if hasattr(frame, 'original_sentence'):
            original_sentence = frame.original_sentence
            text_widget.configure(state="normal")
            
            # Check if we're already showing the original (storing visibility in frame)
            word_visible = getattr(frame, 'word_visible', False)
            
            if not word_visible:
                # Show original sentence
                text_widget.delete("1.0", tk.END)
                text_widget.insert("1.0", original_sentence)
                button.configure(text=get_translation(self.language, "hide"))
                frame.word_visible = True
            else:
                # Show masked sentence
                masked_sentence = self._create_masked_sentence(frame.original_word, original_sentence)
                text_widget.delete("1.0", tk.END)
                text_widget.insert("1.0", masked_sentence)
                button.configure(text=get_translation(self.language, "show"))
                frame.word_visible = False
                
            text_widget.configure(state="disabled")
    
    def export_docx(self):
        """Export sentences to a Word document."""
        if not self.sentence_widgets:
            messagebox.showwarning(
                get_translation(self.language, "warning_title"),
                get_translation(self.language, "no_sentences_warning")
            )
            return
        
        # Get document title from user
        title = simpledialog.askstring(
            get_translation(self.language, "document_title_prompt"),
            get_translation(self.language, "document_title_prompt"),
            initialvalue="Fill in the Blank"
        )
        if not title:
            return
        
        # Ask if user wants to include analysis
        include_analysis = messagebox.askyesno(
            get_translation(self.language, "include_analysis_title"),
            get_translation(self.language, "include_analysis_prompt")
        )
        
        # If including analysis, check for missing analyses
        if include_analysis:
            # Count sentences without analysis
            missing_analyses = [frame for frame in self.sentence_widgets 
                               if not hasattr(frame, 'analysis') or not frame.analysis]
            
            # If there are missing analyses, generate them
            if missing_analyses:
                if not self.api_service.server_connected:
                    messagebox.showwarning(
                        get_translation(self.language, "server_error_title"),
                        get_translation(self.language, "server_connection_guide")
                    )
                    # Continue without analyses
                    include_analysis = False
                elif r'{word}' not in self.api_service.settings_service.get_settings("analysis_prompt") or r'{sentence}' not in self.api_service.settings_service.get_settings("analysis_prompt"):
                    messagebox.showerror(
                        get_translation(self.language, "error_title"),
                        get_translation(self.language, "invalid_prompt_format")
                    )
                    return
                else:
                    # Create progress dialog
                    progress_window = tk.Toplevel(self)
                    progress_window.title(get_translation(self.language, "generating_analyses"))
                    progress_window.geometry("400x100")
                    progress_window.transient(self)
                    progress_window.grab_set()
                    progress_window.resizable(False, False)
                    
                    # Create progress frame
                    progress_frame = ttk.Frame(progress_window, padding="10")
                    progress_frame.pack(fill=tk.BOTH, expand=True)
                    
                    # Labels
                    ttk.Label(progress_frame, text=get_translation(self.language, "generating_analyses")).pack(pady=(0, 5))
                    progress_label = ttk.Label(progress_frame, text=get_translation(self.language, "progress_indicator").format(0, len(missing_analyses)))
                    progress_label.pack(pady=(0, 5))
                    
                    # Progress bar
                    progress_var = tk.DoubleVar()
                    progress_bar = ttk.Progressbar(progress_frame, variable=progress_var, maximum=100)
                    progress_bar.pack(fill=tk.X, pady=(0, 10))
                    
                    # Update progress
                    progress_window.update()
                    
                    prompt = self.api_service.settings_service.get_settings("analysis_prompt")

                    if r'{word}' not in prompt or r'{sentence}' not in prompt:
                        messagebox.showerror(
                            get_translation(self.language, "error_title"),
                            get_translation(self.language, "invalid_prompt_format")
                        )
                        return
                    # Generate analyses one by one
                    for i, frame in enumerate(missing_analyses):
                        word = frame.original_word
                        sentence = frame.original_sentence
                        
                        # Update progress
                        progress_var.set((i / len(missing_analyses)) * 100)
                        progress_label.configure(text=get_translation(self.language, "progress_indicator").format(i+1, len(missing_analyses)))
                        progress_window.update()
                        
                        # Generate analysis
                        formatted_prompt = prompt.format(word=word, sentence=sentence)
                        analysis = self.api_service.generate_sentence(word, formatted_prompt)
                        
                        # Store analysis
                        if analysis:
                            frame.analysis = analysis
                        
                        # Update progress again
                        progress_var.set(((i+1) / len(missing_analyses)) * 100)
                        progress_label.configure(text=get_translation(self.language, "progress_indicator").format(i+1, len(missing_analyses)))
                        progress_window.update()
                    
                    # Close progress window
                    progress_window.destroy()
        
        # Ask for save location
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
            title=get_translation(self.language, "save_document_title"),
            initialfile=f"{title}.docx"
        )
        
        if not file_path:
            return
        
        # Create document
        doc = Document()
        
        # Add title
        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(16)
        
        # Add exercises with blanks first
        for i, frame in enumerate(self.sentence_widgets, 1):
            if frame.winfo_exists():
                # Get the text widget to display the current content
                text_widget = None
                for child in frame.winfo_children():
                    if isinstance(child, tk.Text):
                        text_widget = child
                        break
                
                if text_widget:
                    # Add exercise with blanks (use what's currently displayed)
                    para = doc.add_paragraph()
                    para.add_run(f"{i}. ").bold = True
                    
                    # Get what's currently in the text widget
                    current_text = text_widget.get("1.0", "end-1c")
                    para.add_run(current_text)
        
        # Add page break before answer key
        doc.add_page_break()
        
        # Add answer key title
        answer_title = doc.add_paragraph()
        answer_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        answer_title_run = answer_title.add_run("Answer Key")
        answer_title_run.bold = True
        answer_title_run.font.size = Pt(14)
        
        # Extract and list all blanked words with analysis if available
        for i, frame in enumerate(self.sentence_widgets, 1):
            if frame.winfo_exists():
                # Get the text widget
                text_widget = None
                for child in frame.winfo_children():
                    if isinstance(child, tk.Text):
                        text_widget = child
                        break
                
                if text_widget:
                    # Get the masked text from the text widget
                    masked = text_widget.get("1.0", "end-1c")
                    # Get the original sentence from the frame
                    original = frame.original_sentence
                    
                    # Find all blanks and their answers using improved algorithm
                    words = self._extract_blanked_words_improved(original, masked)
                    
                    if words:
                        para = doc.add_paragraph()
                        para.add_run(f"{i}. ").bold = True
                        
                        # Add word and analysis if available
                        if include_analysis and hasattr(frame, 'analysis') and frame.analysis:
                            para.add_run(f"{words[0]}; [{frame.analysis}]")
                        else:
                            para.add_run(", ".join(words))
                    else:
                        # Fallback to using the original word if no blanks found
                        if hasattr(frame, 'original_word'):
                            para = doc.add_paragraph()
                            para.add_run(f"{i}. ").bold = True
                            if include_analysis and hasattr(frame, 'analysis') and frame.analysis:
                                para.add_run(f"{frame.original_word}; [{frame.analysis}]")
                            else:
                                para.add_run(frame.original_word)
        
        # Save the document
        try:
            doc.save(file_path)
            messagebox.showinfo(
                get_translation(self.language, "export_success_title"),
                get_translation(self.language, "export_success_msg")
            )
        except Exception as e:
            messagebox.showerror(
                get_translation(self.language, "error_title"),
                str(e)
            )

    def _extract_blanked_words_improved(self, original, masked):
        """
        Extract words that have been blanked out in the masked sentence.
        More accurate than the previous method.
        """
        blank_answers = []
        
        # If there's no underscore in the masked sentence, there are no blanks
        if '_' not in masked:
            return []
        
        # First, identify indices where blanks appear
        orig_words = original.split()
        mask_words = masked.split()
        
        # Make sure the sentences have the same word count
        if len(orig_words) != len(mask_words):
            # Different word count indicates mismatch, fall back to advanced parsing
            return self._extract_blanks_by_pattern(original, masked)
        
        # Compare words at the same positions
        for i, (orig_word, mask_word) in enumerate(zip(orig_words, mask_words)):
            clean_orig = orig_word.strip('.,!?;:"\'()')
            clean_mask = mask_word.strip('.,!?;:"\'()')
            
            # Check if this is a blanked word (has underscores)
            if '_' in clean_mask:
                # Extract the original word without punctuation
                if clean_orig not in blank_answers:
                    blank_answers.append(clean_orig)
        
        # If we didn't find any blanks with the direct approach, try advanced pattern matching
        if not blank_answers:
            return self._extract_blanks_by_pattern(original, masked)
        
        return blank_answers

    def _extract_blanks_by_pattern(self, original, masked):
        """Extract blanks using pattern matching for more complex cases."""
        blank_answers = []
        
        # Find all masks (patterns of first letter followed by underscores)
        mask_patterns = []
        for word in masked.split():
            clean_word = word.strip('.,!?;:"\'()')
            if len(clean_word) > 1 and '_' in clean_word:
                # Get the pattern of the masked word
                mask_patterns.append((clean_word, len(clean_word), clean_word[0].lower()))
        
        # For each mask pattern, find the corresponding word in the original
        for mask, length, first_letter in mask_patterns:
            for word in original.split():
                clean_word = word.strip('.,!?;:"\'()')
                if (len(clean_word) == length and
                    clean_word[0].lower() == first_letter and
                    clean_word.lower() not in [w.lower() for w in blank_answers]):
                    blank_answers.append(clean_word)
                    break
        
        return blank_answers

    def show_all_words(self):
        """Show or hide all words in all sentences."""
        try:
            if not self.sentence_widgets:
                return
                
            # Determine the action based on any visible word
            show_all = not any(hasattr(frame, 'word_visible') and frame.word_visible 
                              for frame in self.sentence_widgets)
            
            # Update button text
            self.show_all_btn.configure(
                text=get_translation(self.language, "hide_all" if show_all else "show_all")
            )
            
            for frame in self.sentence_widgets:
                if frame.winfo_exists():
                    # Find the text widget and show button
                    text_widget = None
                    show_button = None
                    
                    for child in frame.winfo_children():
                        if isinstance(child, tk.Text):
                            text_widget = child
                        elif isinstance(child, ttk.Frame):  # Buttons frame
                            for button in child.winfo_children():
                                if "show_button" in str(button):
                                    show_button = button
                                    break
                    
                    if text_widget and show_button:
                        text_widget.configure(state="normal")
                        
                        if show_all:
                            # Show original sentence
                            text_widget.delete("1.0", tk.END)
                            text_widget.insert("1.0", frame.original_sentence)
                            show_button.configure(
                                text=get_translation(self.language, "hide")
                            )
                            frame.word_visible = True
                        else:
                            # Show masked sentence
                            masked_sentence = self._create_masked_sentence(frame.original_word, frame.original_sentence)
                            text_widget.delete("1.0", tk.END)
                            text_widget.insert("1.0", masked_sentence)
                            show_button.configure(
                                text=get_translation(self.language, "show")
                            )
                            frame.word_visible = False
                        
                        text_widget.configure(state="disabled")
        except Exception as e:
            print(f"Error in show_all_words: {e}")
            # If we're running in a development environment, re-raise the error for debugging
            import os
            if os.environ.get('LEXIGEN_DEBUG'):
                raise

    def clear_sentences(self):
        """Delete all sentences."""
        # First destroy all sentence widgets
        for widget in self.sentence_widgets:
            widget.destroy()
        self.sentence_widgets.clear()
        self._update_buttons_state()
        
        # Reset canvas and container
        self.canvas.delete("all")  # Delete all canvas items
        
        # Recreate the sentences container
        self.sentences_container.destroy()
        self.sentences_container = ttk.Frame(self.canvas)
        self.canvas_frame = self.canvas.create_window((0, 0), window=self.sentences_container, anchor="nw")
        
        # Rebind events
        self.sentences_container.bind('<Configure>', self._on_frame_configure)
        
        # Reset canvas scroll region
        self.canvas.configure(scrollregion=(0, 0, 0, 0))
        
        # Force update to clear any remaining visual artifacts
        self.update_idletasks()
        self.canvas.update()
        
        # Notify about sentence change
        if self.on_sentences_changed:
            self.on_sentences_changed(False)

    def _copy_sentence(self, text_widget):
        """Copy sentence to clipboard."""
        # Find the parent frame
        frame = text_widget.master
        
        # Get the appropriate sentence based on visibility
        if hasattr(frame, 'word_visible') and frame.word_visible:
            # Copy the original sentence with filled blanks
            sentence = frame.original_sentence
        else:
            # Copy the masked sentence with blanks (what's currently displayed)
            sentence = text_widget.get("1.0", "end-1c")
        
        # Strip any trailing newlines or spaces
        sentence = sentence.strip()
        
        # Copy to clipboard
        self.clipboard_clear()
        self.clipboard_append(sentence)
        
        # Show brief visual feedback on the copy button
        for child in frame.winfo_children():
            if isinstance(child, ttk.Frame):  # Buttons frame
                for button in child.winfo_children():
                    if "copy_button" in str(button):
                        original_text = button.cget("text")
                        button.configure(text=get_translation(self.language, "checkmark"))
                        
                        # Reset the button text after a short delay
                        def reset_text():
                            if button.winfo_exists():
                                button.configure(text=original_text)
                        
                        # Schedule reset after 1 second
                        self.after(1000, reset_text)
                        break

    def _update_buttons_state(self):
        has_sentences = len(self.sentence_widgets) > 0
        
        # Menu button should always be enabled to allow loading history
        if hasattr(self, 'menu_btn'):
            self.menu_btn.configure(state="normal")
        
        # Show all and delete buttons are only enabled when there are sentences
        sentence_dependent_state = "normal" if has_sentences else "disabled"
        if hasattr(self, 'show_all_btn'):
            self.show_all_btn.configure(state=sentence_dependent_state)
            
        if hasattr(self, 'delete_btn'):
            self.delete_btn.configure(state=sentence_dependent_state)
        
        if self.on_sentences_changed:
            self.on_sentences_changed(has_sentences)

    def _delete_sentence(self, frame):
        """Delete a single sentence widget."""
        if frame in self.sentence_widgets:
            self.sentence_widgets.remove(frame)
        frame.destroy()
        self._update_buttons_state()
        
        # Renumber remaining sentences
        for i, widget in enumerate(self.sentence_widgets):
            widget.order_label.configure(text=f"{i + 1}.")
        
        # Notify about sentence change
        if self.on_sentences_changed:
            self.on_sentences_changed(len(self.sentence_widgets) > 0)
        
        # If this was the last sentence, reset the canvas completely
        if not self.sentence_widgets:
            # Reset canvas and container similar to clear_sentences
            self.canvas.delete("all")  # Delete all canvas items
            
            # Recreate the sentences container
            self.sentences_container.destroy()
            self.sentences_container = ttk.Frame(self.canvas)
            self.canvas_frame = self.canvas.create_window((0, 0), window=self.sentences_container, anchor="nw")
            
            # Rebind events
            self.sentences_container.bind('<Configure>', self._on_frame_configure)
            
            # Reset canvas scroll region
            self.canvas.configure(scrollregion=(0, 0, 0, 0))
            
            # Force update to clear any remaining visual artifacts
            self.update_idletasks()
            self.canvas.update()

    def _regenerate_sentence(self, word, frame):
        """Regenerate the sentence for a specific word."""
        # Get prompt from settings
        prompt = self.main_window.settings_service.get_settings("generation_prompt")
        if not prompt:
            prompt = DEFAULT_CONFIG["generation_prompt"]
        
        # Disable the regenerate button and change it to progress indicator
        regen_btn = None
        for child in frame.winfo_children():
            if isinstance(child, ttk.Frame):  # This is the buttons frame
                for button in child.winfo_children():
                    if "regen_button" in str(button):
                        regen_btn = button
                        break
        
        if regen_btn:
            original_text = regen_btn.cget("text")
            regen_btn.config(text="...", state="disabled")
        
        # Check if we have a context attachment
        if hasattr(self.main_window, 'context') and self.main_window.context:
            context_attachment_prompt = self.main_window.settings_service.get_settings("context_attachment_prompt")
            prompt = context_attachment_prompt.format(context=self.main_window.context) + "\n" + prompt

        # Generate a new sentence with retry logic
        max_attempts = 3
        for attempt in range(max_attempts):
            # Get the new sentence
            sentence = self.api_service.generate_sentence(word, prompt)
            
            if sentence:
                # Store the new original sentence in the frame
                frame.original_sentence = sentence
                
                # Create masked sentence
                masked_sentence = self._create_masked_sentence(word, sentence)
                
                # Find the text widget in this frame
                text_widget = None
                for child in frame.winfo_children():
                    if isinstance(child, tk.Text):
                        text_widget = child
                        break
                
                # Get the show button
                show_btn = None
                for child in frame.winfo_children():
                    if isinstance(child, ttk.Frame):  # This is the buttons frame
                        for button in child.winfo_children():
                            if "show_button" in str(button):
                                show_btn = button
                                break
                
                if text_widget:
                    # Update the text widget
                    text_widget.config(state="normal")
                    text_widget.delete("1.0", "end")
                    text_widget.insert("1.0", masked_sentence)
                    text_widget.config(state="disabled")
                    
                    # Reset the button
                    if regen_btn:
                        regen_btn.config(text=original_text, state="normal")
                    
                    # Reset the show button text
                    if show_btn:
                        show_btn.config(text=get_translation(self.language, "show"))
                    
                    # Clear any existing analysis since it's no longer valid
                    if hasattr(frame, 'analysis'):
                        delattr(frame, 'analysis')
                    
                    # Adjust height
                    text_widget.after(10, lambda tw=text_widget: self._adjust_text_height(tw))
                    
                    return True
            
            # If we've reached the maximum attempts, reset the button
            if attempt == max_attempts - 1:
                if regen_btn:
                    regen_btn.config(text=original_text, state="normal")
                messagebox.showerror(
                    get_translation(self.language, "error_title"),
                    get_translation(self.language, "sentence_generation_failed")
                )
                return False
                
        # If we're here, generation failed
        if regen_btn:
            regen_btn.config(text=original_text, state="normal")
        return False
    
    def _adjust_text_height(self, text_widget):
        """Adjust the height of a text widget based on its content."""
        text_widget.configure(state="normal")
        text_widget.see("end")
        num_lines = text_widget.count("1.0", "end", "displaylines")[0]
        text_widget.configure(height=max(2, num_lines))
        text_widget.configure(state="disabled")

    def update_texts(self, language):
        """Update all texts in the widget after language change."""
        self.language = language
        
        # Update the frame title
        self.configure(text=get_translation(language, "generated_sentences"))
        
        # Update control buttons
        if hasattr(self, 'menu_btn'):
            self.menu_btn.configure(text=get_translation(language, "menu_button_main"))
        
        # Update show/hide all button based on its current state
        if hasattr(self, 'show_all_btn'):
            current_text = self.show_all_btn.cget("text")
            if current_text == get_translation(self.language, "hide_all") or "hide" in current_text.lower() or "隐藏" in current_text:
                self.show_all_btn.configure(text=get_translation(language, "hide_all"))
            else:
                self.show_all_btn.configure(text=get_translation(language, "show_all"))
        
        # Update delete button
        if hasattr(self, 'delete_btn'):
            self.delete_btn.configure(text=get_translation(language, "delete_all"))
        
        # Update individual sentence widgets
        for frame in self.sentence_widgets:
            if not frame.winfo_exists():
                continue
            
            # Find the text widget and buttons
            text_widget = None
            show_button = None
            copy_button = None
            regen_button = None
            menu_button = None
            
            for child in frame.winfo_children():
                if isinstance(child, tk.Text):
                    text_widget = child
                elif isinstance(child, ttk.Frame):  # Buttons frame
                    for button in child.winfo_children():
                        if "show_button" in str(button):
                            show_button = button
                        elif "copy_button" in str(button):
                            copy_button = button
                        elif "regen_button" in str(button):
                            regen_button = button
                        elif "menu_button" in str(button):
                            menu_button = button
            
            # Update show/hide button based on current visibility state
            if show_button:
                if hasattr(frame, 'word_visible') and frame.word_visible:
                    show_button.configure(text=get_translation(language, "hide"))
                else:
                    show_button.configure(text=get_translation(language, "show"))
            
            # Update other buttons
            if copy_button:
                copy_button.configure(text=get_translation(language, "copy"))
            if regen_button:
                regen_button.configure(text=get_translation(language, "regenerate_button"))
            if menu_button:
                menu_button.configure(text=get_translation(language, "menu_button"))

    def _show_main_menu(self):
        """Show the main menu for the sentence frame."""
        # Create menu
        menu = tk.Menu(self, tearoff=0)
        
        # Add Export to Word option (disabled if no sentences)
        has_sentences = len(self.sentence_widgets) > 0
        menu.add_command(
            label=get_translation(self.language, "export_docx"),
            command=self.export_docx,
            state="normal" if has_sentences else "disabled"
        )
        
        # Add Save History option (disabled if no sentences)
        menu.add_command(
            label=f"{get_translation(self.language, 'save_history')}",
            command=self.save_history,
            state="normal" if has_sentences else "disabled"
        )
        
        # Add Load History option (always enabled)
        menu.add_command(
            label=f"{get_translation(self.language, 'load_history')}",
            command=self.load_history
        )
        
        # Get button position
        btn = self.menu_btn
        x = btn.winfo_rootx()
        y = btn.winfo_rooty() + btn.winfo_height()
        
        # Show menu and bind to close on click outside
        menu.post(x, y)
        menu.bind("<Unmap>", lambda e: menu.destroy())
        
        # Bind to close on any click
        def close_menu(e):
            menu.destroy()
            self.unbind("<Button-1>")
        
        self.bind("<Button-1>", close_menu)
    
    def save_history(self):
        """Save the history of sentences to a YAML file."""
        if not self.sentence_widgets:
            messagebox.showwarning(
                get_translation(self.language, "warning_title"),
                get_translation(self.language, "no_sentences_to_save")
            )
            return
        
        # Ask user where to save the file
        file_path = filedialog.asksaveasfilename(
            title=get_translation(self.language, "save_history_title"),
            defaultextension=".yaml",
            filetypes=[("YAML files", "*.yaml"), ("All files", "*.*")]
        )
        
        if not file_path:
            return  # User cancelled
        
        # Prepare data structure to save
        history_data = {
            "sentences": [],
            "context": None
        }
        
        # Add context if available
        if hasattr(self.main_window, "context") and self.main_window.context:
            history_data["context"] = self.main_window.context
        
        # Collect all sentence data
        for frame in self.sentence_widgets:
            # Get original word
            original_word = frame.original_word
            
            # Get sentence
            text_widget = None
            for child in frame.winfo_children():
                if isinstance(child, tk.Text):
                    text_widget = child
                    break
            
            if text_widget:
                # Get the visible (masked) sentence
                masked_sentence = text_widget.get("1.0", "end-1c")
                
                # Get original sentence (if available, otherwise use masked)
                original_sentence = getattr(frame, "original_sentence", masked_sentence)
                
                # Get analysis if available
                analysis = getattr(frame, "analysis", None)
                
                # Add to history
                sentence_data = {
                    "word": original_word,
                    "sentence": original_sentence,
                    "masked_sentence": masked_sentence,
                    "analysis": analysis
                }
                
                history_data["sentences"].append(sentence_data)
        
        try:
            # Save to YAML
            with open(file_path, 'w', encoding='utf-8') as f:
                yaml.safe_dump(history_data, f, allow_unicode=True, default_flow_style=False)
            
            messagebox.showinfo(
                get_translation(self.language, "export_success_title"),
                get_translation(self.language, "history_save_success")
            )
        except Exception as e:
            messagebox.showerror(
                get_translation(self.language, "error_title"),
                str(e)
            )
    
    def load_history(self):
        """Load sentence history from a YAML file."""
        # Ask user for file to load
        file_path = filedialog.askopenfilename(
            title=get_translation(self.language, "load_history_title"),
            filetypes=[("YAML files", "*.yaml"), ("All files", "*.*")]
        )
        
        if not file_path:
            return  # User cancelled
        
        try:
            # Load from YAML
            with open(file_path, 'r', encoding='utf-8') as f:
                history_data = yaml.safe_load(f) or {}
            
            if not history_data or "sentences" not in history_data:
                raise ValueError("Invalid history file format")
            
            # Clear existing sentences
            self.clear_sentences()
            
            # Load context if available
            if "context" in history_data and history_data["context"] and hasattr(self.main_window, "context"):
                self.main_window.context = history_data["context"]
            
            # Load all sentences
            for sentence_data in history_data["sentences"]:
                word = sentence_data.get("word", "")
                sentence = sentence_data.get("sentence", "")
                
                if word and sentence:
                    # Add the sentence to UI
                    self.add_sentence(word, sentence)
                    
                    # If this is the last sentence widget
                    if self.sentence_widgets:
                        frame = self.sentence_widgets[-1]
                        
                        # Store original sentence
                        frame.original_sentence = sentence
                        
                        # Store analysis if available
                        if "analysis" in sentence_data and sentence_data["analysis"]:
                            frame.analysis = sentence_data["analysis"]
            
            messagebox.showinfo(
                get_translation(self.language, "export_success_title"),
                get_translation(self.language, "history_load_success")
            )
            
            # Update button states
            self._update_buttons_state()
            
        except Exception as e:
            messagebox.showerror(
                get_translation(self.language, "error_title"),
                get_translation(self.language, "history_load_error").format(error=str(e))
            )

    def _show_menu(self, frame):
        """Show the menu for a sentence frame."""
        # Create menu
        menu = tk.Menu(self, tearoff=0)
        
        # Add menu items
        menu.add_command(
            label=get_translation(self.language, "move_up"),
            command=lambda: self._move_sentence(frame, -1)
        )
        menu.add_command(
            label=get_translation(self.language, "move_down"),
            command=lambda: self._move_sentence(frame, 1)
        )
        menu.add_command(
            label=get_translation(self.language, "analyze"),
            command=lambda: self._show_analysis(frame)
        )
        menu.add_command(
            label=get_translation(self.language, "edit"),
            command=lambda: self._edit_sentence(frame)
        )
        
        # Add Designate Tense submenu
        tense_menu = tk.Menu(menu, tearoff=0)
        menu.add_cascade(label=get_translation(self.language, "designate_tense"), menu=tense_menu)
        
        # Add tense options
        present_menu = tk.Menu(tense_menu, tearoff=0)
        tense_menu.add_cascade(label=get_translation(self.language, "present"), menu=present_menu)
        present_menu.add_command(
            label=get_translation(self.language, "simple"),
            command=lambda: self._generate_with_tense(frame, "Present Simple")
        )
        present_menu.add_command(
            label=get_translation(self.language, "continuous"), 
            command=lambda: self._generate_with_tense(frame, "Present Continuous")
        )
        present_menu.add_command(
            label=get_translation(self.language, "perfect"), 
            command=lambda: self._generate_with_tense(frame, "Present Perfect")
        )
        present_menu.add_command(
            label=get_translation(self.language, "perfect_continuous"), 
            command=lambda: self._generate_with_tense(frame, "Present Perfect Continuous")
        )
        
        past_menu = tk.Menu(tense_menu, tearoff=0)
        tense_menu.add_cascade(label=get_translation(self.language, "past"), menu=past_menu)
        past_menu.add_command(
            label=get_translation(self.language, "simple"), 
            command=lambda: self._generate_with_tense(frame, "Past Simple")
        )
        past_menu.add_command(
            label=get_translation(self.language, "continuous"), 
            command=lambda: self._generate_with_tense(frame, "Past Continuous")
        )
        past_menu.add_command(
            label=get_translation(self.language, "perfect"), 
            command=lambda: self._generate_with_tense(frame, "Past Perfect")
        )
        past_menu.add_command(
            label=get_translation(self.language, "perfect_continuous"), 
            command=lambda: self._generate_with_tense(frame, "Past Perfect Continuous")
        )
        
        future_menu = tk.Menu(tense_menu, tearoff=0)
        tense_menu.add_cascade(label=get_translation(self.language, "future"), menu=future_menu)
        future_menu.add_command(
            label=get_translation(self.language, "simple"), 
            command=lambda: self._generate_with_tense(frame, "Future Simple")
        )
        future_menu.add_command(
            label=get_translation(self.language, "continuous"), 
            command=lambda: self._generate_with_tense(frame, "Future Continuous")
        )
        future_menu.add_command(
            label=get_translation(self.language, "perfect"), 
            command=lambda: self._generate_with_tense(frame, "Future Perfect")
        )
        future_menu.add_command(
            label=get_translation(self.language, "perfect_continuous"), 
            command=lambda: self._generate_with_tense(frame, "Future Perfect Continuous")
        )
        
        past_future_menu = tk.Menu(tense_menu, tearoff=0)
        tense_menu.add_cascade(label=get_translation(self.language, "past_future"), menu=past_future_menu)
        past_future_menu.add_command(
            label=get_translation(self.language, "simple"), 
            command=lambda: self._generate_with_tense(frame, "Past Future Simple")
        )
        past_future_menu.add_command(
            label=get_translation(self.language, "continuous"), 
            command=lambda: self._generate_with_tense(frame, "Past Future Continuous")
        )
        past_future_menu.add_command(
            label=get_translation(self.language, "perfect"), 
            command=lambda: self._generate_with_tense(frame, "Past Future Perfect")
        )
        past_future_menu.add_command(
            label=get_translation(self.language, "perfect_continuous"), 
            command=lambda: self._generate_with_tense(frame, "Past Future Perfect Continuous")
        )
        
        # Subjunctive Mood submenu
        subjunctive_menu = tk.Menu(tense_menu, tearoff=0)
        tense_menu.add_cascade(label=get_translation(self.language, "subjunctive_mood"), menu=subjunctive_menu)
        subjunctive_menu.add_command(
            label=get_translation(self.language, "present"),
            command=lambda: self._generate_with_tense(frame, "Present Subjunctive Mood")
        )
        subjunctive_menu.add_command(
            label=get_translation(self.language, "past"),
            command=lambda: self._generate_with_tense(frame, "Past Subjunctive Mood")
        )
        
        # Conditional submenu
        conditional_menu = tk.Menu(tense_menu, tearoff=0)
        tense_menu.add_cascade(label=get_translation(self.language, "conditional"), menu=conditional_menu)
        conditional_menu.add_command(
            label=get_translation(self.language, "zero_conditional"),
            command=lambda: self._generate_with_tense(frame, "Zero Conditional")
        )
        conditional_menu.add_command(
            label=get_translation(self.language, "first_conditional"),
            command=lambda: self._generate_with_tense(frame, "First Conditional")
        )
        conditional_menu.add_command(
            label=get_translation(self.language, "second_conditional"),
            command=lambda: self._generate_with_tense(frame, "Second Conditional")
        )
        conditional_menu.add_command(
            label=get_translation(self.language, "third_conditional"),
            command=lambda: self._generate_with_tense(frame, "Third Conditional")
        )

        # Imperative Mood submenu
        tense_menu.add_command(
            label=get_translation(self.language, "imperative_mood"),
            command=lambda: self._generate_with_tense(frame, "Imperative Mood")
        )
        
        menu.add_separator()
        menu.add_command(
            label=get_translation(self.language, "delete"),
            command=lambda: self._delete_sentence(frame)
        )
        
        # Get button position
        btn = frame.menu_btn
        x = btn.winfo_rootx()
        y = btn.winfo_rooty() + btn.winfo_height()
        
        # Show menu and bind to close on click outside
        menu.post(x, y)
        menu.bind("<Unmap>", lambda e: menu.destroy())
        
        # Bind to close on any click
        def close_menu(e):
            menu.destroy()
            self.unbind("<Button-1>")
        
        self.bind("<Button-1>", close_menu)

    def _move_sentence(self, frame, direction):
        index = self.sentence_widgets.index(frame)
        new_index = index + direction
        
        if 0 <= new_index < len(self.sentence_widgets):
            # Remove from current position
            self.sentence_widgets.pop(index)
            # Insert at new position
            self.sentence_widgets.insert(new_index, frame)
            
            # Update grid positions
            for i, widget in enumerate(self.sentence_widgets):
                widget.grid(row=i)
            
            # Update all order numbers
            for i, widget in enumerate(self.sentence_widgets):
                widget.order_label.configure(text=f"{i + 1}.")
            
            # Notify about sentence change
            if self.on_sentences_changed:
                self.on_sentences_changed(len(self.sentence_widgets) > 0)

    def _show_analysis(self, frame):
        """Show analysis window for the sentence."""
        # Extract the original word and sentence from the frame
        word = frame.original_word
        sentence = frame.original_sentence
        
        # Find the text widget
        text_widget = None
        for child in frame.winfo_children():
            if isinstance(child, tk.Text):
                text_widget = child
                break
                
        if not text_widget:
            return
        
        # Create and show analysis window
        analysis_window = AnalysisWindow(self, word, sentence, self.api_service, self.language, text_widget)
        
        # Set the frame reference for storing analysis
        analysis_window.sentence_frame = frame
        
        # If we already have an analysis, display it
        if hasattr(frame, 'analysis') and frame.analysis:
            analysis_window.load_existing_analysis(frame.analysis)
            
    def _edit_sentence(self, frame):
        """Edit a sentence."""
        word = frame.original_word
        sentence = frame.original_sentence
        
        # Create and show edit window
        EditSentenceWindow(self, word, sentence, self.api_service, self.language, frame)
    
    def _generate_with_tense(self, frame, tense):
        """Generate a sentence with a specified tense."""
        if not self.api_service.server_connected:
            messagebox.showwarning(
                get_translation(self.language, "server_error_title"),
                get_translation(self.language, "server_connection_guide")
            )
            return
        
        word = frame.original_word
        frame.tense = tense
        
        # Get tense prompt template from settings
        prompt_template = self.api_service.settings_service.get_settings("tense_prompt")
        try:
            prompt = prompt_template.format(word=word, tense=tense)
        except Exception as e:
            messagebox.showerror(
                get_translation(self.language, "error_title"),
                get_translation(self.language, "invalid_prompt_format").format(error=str(e))
            )
            return
        
        # Generate new sentence with the specified tense
        new_sentence = self.api_service.generate_sentence(word, prompt)
        
        if new_sentence:
            # Update original sentence in frame
            frame.original_sentence = new_sentence
            
            # Create masked sentence
            masked_sentence = self._create_masked_sentence(word, new_sentence)
            
            # Find the text widget
            text_widget = None
            for child in frame.winfo_children():
                if isinstance(child, tk.Text):
                    text_widget = child
                    break
                    
            if text_widget:
                # Update the text widget
                text_widget.configure(state="normal")
                text_widget.delete("1.0", tk.END)
                text_widget.insert("1.0", masked_sentence)
                text_widget.configure(state="disabled")
                
                # Remove analysis if it exists as the sentence has changed
                if hasattr(frame, 'analysis'):
                    delattr(frame, 'analysis')
                
                # Adjust height
                self._adjust_text_height(text_widget)

class AnalysisWindow(tk.Toplevel):
    def __init__(self, parent, word, sentence, api_service, language, text_widget):
        super().__init__(parent)
        self.parent = parent
        self.word = word
        self.sentence = sentence
        self.api_service = api_service
        self.language = language
        self.text_widget = text_widget
        self.sentence_frame = None  # Will be set by caller
        self.title(get_translation(self.language, "word_analysis"))

        # Set geometry and make modal
        width = 600
        height = 300
        x = (self.winfo_screenwidth() // 2) - (width // 2)
        y = (self.winfo_screenheight() // 2) - (height // 2)
        self.geometry(f"{width}x{height}+{x}+{y}")
        self.transient(parent)
        self.grab_set()
        
        # Setup UI
        self.setup_ui()
        
    def setup_ui(self):
        # Main container
        main_frame = ttk.Frame(self, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Word and sentence display
        info_frame = ttk.Frame(main_frame)
        info_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(info_frame, text=f"{get_translation(self.language, 'word')}: ").pack(side=tk.LEFT)
        ttk.Label(info_frame, text=self.word, font=("", 10, "bold")).pack(side=tk.LEFT)
        
        sentence_frame = ttk.Frame(main_frame)
        sentence_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(sentence_frame, text=f"{get_translation(self.language, 'sentence')}: ").pack(anchor=tk.W)
        sentence_text = tk.Text(sentence_frame, wrap=tk.WORD, height=3)
        sentence_text.insert("1.0", self.sentence)
        sentence_text.configure(state="disabled")
        sentence_text.pack(fill=tk.X)
        
        # Analysis section
        analysis_frame = ttk.Frame(main_frame)
        analysis_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        
        ttk.Label(analysis_frame, text=f"{get_translation(self.language, 'analysis')}: ").pack(anchor=tk.W)
        self.analysis_text = scrolledtext.ScrolledText(analysis_frame, wrap=tk.WORD, height=5)
        self.analysis_text.pack(fill=tk.BOTH, expand=True)
        
        # Buttons
        buttons_frame = ttk.Frame(main_frame)
        buttons_frame.pack(fill=tk.X)
        
        ttk.Button(buttons_frame, text=get_translation(self.language, "regenerate_analysis"), 
                  command=self._regenerate_analysis).pack(side=tk.LEFT, padx=(0, 5))
        
        # Add an Edit/Save toggle button
        self.edit_btn = ttk.Button(buttons_frame, text=get_translation(self.language, "edit"), 
                                  command=self._toggle_edit_mode)
        self.edit_btn.pack(side=tk.LEFT, padx=(0, 5))
        
        ttk.Button(buttons_frame, text=get_translation(self.language, "close"), 
                  command=self._on_close).pack(side=tk.RIGHT)
        
        # Start analysis (if not already available)
        self._generate_analysis()
        
    def load_existing_analysis(self, analysis):
        """Load an existing analysis"""
        self.analysis_result = analysis
        self._display_analysis()

    def _toggle_edit_mode(self):
        """Toggle between edit and read mode for analysis"""
        if self.analysis_text.cget("state") == "normal":
            # Save the analysis when switching from edit mode
            self.analysis_result = self.analysis_text.get("1.0", tk.END).strip()
            
            # Store the analysis in the sentence frame
            if self.sentence_frame is not None:
                self.sentence_frame.analysis = self.analysis_result
                
            self.analysis_text.configure(state="disabled")
            self.edit_btn.configure(text=get_translation(self.language, "edit"))
        else:
            # Enable editing
            self.analysis_text.configure(state="normal")
            self.edit_btn.configure(text=get_translation(self.language, "save"))

    def _display_analysis(self):
        """Display the analysis results."""
        if hasattr(self, 'analysis_result') and self.analysis_result:
            self.analysis_text.configure(state="normal")
            self.analysis_text.delete("1.0", tk.END)
            self.analysis_text.insert("1.0", self.analysis_result)
            self.analysis_text.configure(state="disabled")

    def _generate_analysis(self):
        """Generate analysis for the given word and sentence."""
        # Check if we already have an analysis stored in the frame
        if self.sentence_frame is not None and hasattr(self.sentence_frame, 'analysis') and self.sentence_frame.analysis:
            self.analysis_result = self.sentence_frame.analysis
            self._display_analysis()
            return
            
        if not self.api_service.server_connected:
            self.analysis_text.configure(state="normal")
            self.analysis_text.delete("1.0", tk.END)
            self.analysis_text.insert("1.0", get_translation(self.language, "server_not_connected"))
            self.analysis_text.configure(state="disabled")
            return

        self.analysis_text.configure(state="normal")
        self.analysis_text.delete("1.0", tk.END)
        self.analysis_text.insert("1.0", get_translation(self.language, "analyzing"))
        self.analysis_text.configure(state="disabled")
        self.update()

        # Get analysis prompt
        analysis_prompt = self.api_service.settings_service.get_settings("analysis_prompt")
        
        # Format the prompt with the word and sentence
        prompt = analysis_prompt.format(word=self.word, sentence=self.sentence)
        
        # Use the generation method but with analysis prompt
        self.analysis_result = self._get_analysis(prompt)
        
        # Store the analysis in the sentence frame
        if self.sentence_frame is not None and self.analysis_result:
            self.sentence_frame.analysis = self.analysis_result
        
        # Display the result
        self._display_analysis()
        
    def _regenerate_analysis(self):
        """Force regeneration of analysis."""
        # Clear existing analysis if any
        if self.sentence_frame is not None and hasattr(self.sentence_frame, 'analysis'):
            delattr(self.sentence_frame, 'analysis')
        
        # Generate new analysis
        self._generate_analysis()

    def _on_close(self):
        """Handle window close event."""
        if self.analysis_text.cget("state") == "normal":
            # If in edit mode, save the analysis before closing
            self.analysis_result = self.analysis_text.get("1.0", tk.END).strip()
            
            # Store the analysis in the sentence frame
            if self.sentence_frame is not None:
                self.sentence_frame.analysis = self.analysis_result
        
        self.destroy()

    def _get_analysis(self, prompt):
        """Get analysis from either local model or API."""
        try:
            if self.api_service.using_local_model and self.api_service.local_model is not None:
                # Use local model for analysis
                output = self.api_service.local_model(
                    prompt,
                    max_tokens=256,
                    stop=["</s>", "\n\n"],
                    echo=False
                )
                return output['choices'][0]['text'].strip()
            else:
                # Use remote API
                response = requests.post(
                    self.api_service.api_url,
                    json={
                        "model": self.api_service.model,
                        "prompt": prompt,
                        "stream": False
                    }
                )
                response.raise_for_status()
                result = response.json()
                return result["response"].strip()
        except Exception as e:
            return f"{get_translation(self.language, 'analysis_error')}: {str(e)}"

class EditSentenceWindow(tk.Toplevel):
    def __init__(self, parent, word, sentence, api_service, language, frame):
        super().__init__(parent)
        self.title(get_translation(language, "edit_sentence"))
        width = 600
        height = 220
        x = (self.winfo_screenwidth() // 2) - (width // 2)
        y = (self.winfo_screenheight() // 2) - (height // 2)
        self.geometry(f"{width}x{height}+{x}+{y}")
        self.resizable(True, True)
        
        # Make window modal
        self.transient(parent)
        self.grab_set()
        
        # Bind to window close event
        self.protocol("WM_DELETE_WINDOW", self.destroy)
        
        # Store references
        self.parent = parent
        self.api_service = api_service
        self.language = language
        self.word = word
        self.original_sentence = sentence
        self.frame = frame
        
        # Create main frame
        main_frame = ttk.Frame(self, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Word display
        word_frame = ttk.Frame(main_frame)
        word_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(word_frame, text=f"{get_translation(language, 'word')}: {word}").pack(side=tk.LEFT)
        
        # Sentence editing
        sentence_frame = ttk.Frame(main_frame)
        sentence_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(sentence_frame, text=f"{get_translation(language, 'sentence')}:").pack(anchor=tk.W)
        
        # Sentence text widget
        self.sentence_text = tk.Text(sentence_frame, wrap=tk.WORD, height=5)
        self.sentence_text.insert("1.0", sentence)
        self.sentence_text.pack(fill=tk.X, expand=True, pady=(5, 0))
        
        # Label with warning
        warning_label = ttk.Label(main_frame, text=get_translation(language, "edit_warning"), foreground="red")
        warning_label.pack(pady=(0, 10))
        
        # Buttons frame
        buttons_frame = ttk.Frame(main_frame)
        buttons_frame.pack(fill=tk.X)
        
        save_btn = ttk.Button(
            buttons_frame,
            text=get_translation(language, "save"),
            command=self._save_changes
        )
        save_btn.pack(side=tk.LEFT)
        
        cancel_btn = ttk.Button(
            buttons_frame,
            text=get_translation(language, "cancel"),
            command=self.destroy
        )
        cancel_btn.pack(side=tk.RIGHT)
    
    def _save_changes(self):
        """Save the edited sentence."""
        new_sentence = self.sentence_text.get("1.0", "end-1c").strip()
        
        # Don't do anything if the sentence is empty or unchanged
        if not new_sentence or new_sentence == self.original_sentence:
            self.destroy()
            return
        
        # Update the original sentence in the frame
        self.frame.original_sentence = new_sentence
        
        # Find the text widget
        text_widget = None
        for child in self.frame.winfo_children():
            if isinstance(child, tk.Text):
                text_widget = child
                break
                
        if text_widget:
            # Create new masked sentence
            masked_sentence = self.parent._create_masked_sentence(self.word, new_sentence)
            
            # Update the display
            text_widget.configure(state="normal")
            text_widget.delete("1.0", tk.END)
            text_widget.insert("1.0", masked_sentence)
            text_widget.configure(state="disabled")
            
            # Remove analysis if it exists as the sentence has changed
            if hasattr(self.frame, 'analysis'):
                delattr(self.frame, 'analysis')
            
            # Adjust text height
            self.parent._adjust_text_height(text_widget)
        
        self.destroy()