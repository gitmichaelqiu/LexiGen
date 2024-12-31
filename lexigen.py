import tkinter as tk
from tkinter import ttk, scrolledtext, messagebox, filedialog, simpledialog
import requests
import json
import webbrowser
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import nltk
from nltk.stem import WordNetLemmatizer
from nltk.corpus import wordnet
import threading
import sys

# Version information
VERSION = "1.2.0"

# Language translations
TRANSLATIONS = {
    "English": {
        "settings": "Settings",
        "api_url": "API URL:",
        "model": "Model:",
        "server_status": "Server Status:",
        "server_status_checking": "Server Status: Checking...",
        "server_status_connected": "Server Status: Connected",
        "server_status_error": "Server Status: Error",
        "server_status_not_connected": "Server Status: Not Connected",
        "check_server": "Check Server",
        "setup_help": "Setup Help",
        "check_updates": "Check for Updates",
        "new_version": "New version available",
        "up_to_date": "Up to date",
        "input_words": "Input Words",
        "enter_words": "Enter words separated by commas",
        "generate": "Generate",
        "append": "Append",
        "generating": "Generating sentences...",
        "generated_sentences": "Generated Sentences",
        "export_docx": "Export Docx",
        "show_all": "Show All",
        "hide_all": "Hide All",
        "delete_all": "Delete All",
        "show": "Show",
        "hide": "Hide",
        "copy": "Copy",
        "language": "Language:",
        "server_status_title": "Server Status",
        "server_connected_msg": "Successfully connected to Ollama server!\n\nServer version: {version}",
        "server_error_msg": "Failed to connect to Ollama server.\n\nPlease check if the server is running and the URL is correct.",
        "server_connection_error_msg": "Cannot connect to Ollama server. Please ensure:\n\n1. Ollama is installed\n2. Server is running ('ollama serve')\n3. API URL is correct\n\nError: {error}",
        "server_error_title": "Server Error",
        "server_connection_guide": "Cannot connect to Ollama server. Please ensure:\n1. Ollama is installed\n2. Server is running ('ollama serve')\n3. Model is installed ('ollama pull llama2')\nClick 'Setup Help' for more information.",
        "warning_title": "Warning",
        "no_sentences_warning": "No sentences to export!",
        "document_title_prompt": "Enter the title for your document:",
        "save_document_title": "Save Document As",
        "export_success_title": "Success",
        "export_success_msg": "Document exported successfully!",
        "update_available_title": "Update Available",
        "update_available_msg": "A new version (v{version}) is available!\n\nCurrent version: v{current_version}\n\nWould you like to visit the download page?",
        "no_updates_title": "No Updates",
        "no_updates_msg": "You are using the latest version (v{version})",
        "error_title": "Error",
        "update_check_error": "Failed to check for updates. Please try again later.",
        "update_error_msg": "Failed to check for updates:\n{error}",
        "prompt_error_msg": "Failed to read prompt file: {error}\nUsing default prompt.",
        "generation_error_msg": "Failed to generate sentence: {error}\nPlease check your server settings and connection.",
        "unexpected_error_msg": "An unexpected error occurred: {error}",
        "startup_error_msg": "Error starting application:\n{error}",
        "using_custom_prompt": "Using custom prompt",
        "using_default_prompt": "Using default prompt",
        "toggle_prompt": "Toggle Prompt"
    }
}

def load_translations():
    """Load translations from external files."""
    try:
        # Get the application's base directory
        if getattr(sys, 'frozen', False):
            # Running in a bundle (packaged)
            if sys.platform == 'darwin':
                # For macOS, get the parent directory of the .app bundle
                base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(sys.executable))))
            else:
                # For Windows, get the directory containing the .exe
                base_dir = os.path.dirname(sys.executable)
            print(f"Running in packaged mode, base directory: {base_dir}")
        else:
            # Running in a normal Python environment
            base_dir = os.path.dirname(os.path.abspath(__file__))
            print(f"Running in development mode, base directory: {base_dir}")
        
        translations_dir = os.path.join(base_dir, "translations")
        print(f"Looking for translations in: {translations_dir}")
        
        if not os.path.exists(translations_dir):
            try:
                os.makedirs(translations_dir)
                print(f"Created translations directory: {translations_dir}")
            except Exception as e:
                print(f"Failed to create translations directory: {str(e)}")
        
        if os.path.exists(translations_dir):
            print(f"Translations directory exists. Checking permissions...")
            # Check directory permissions
            try:
                test_perms = os.access(translations_dir, os.R_OK)
                print(f"Directory readable: {test_perms}")
            except Exception as e:
                print(f"Failed to check directory permissions: {str(e)}")
            
            # List and load translation files
            try:
                files = os.listdir(translations_dir)
                print(f"Found files in translations directory: {files}")
                
                for file in files:
                    if file.endswith('.json'):
                        file_path = os.path.join(translations_dir, file)
                        try:
                            print(f"Attempting to read: {file_path}")
                            with open(file_path, 'r', encoding='utf-8') as f:
                                lang_translations = json.load(f)
                                TRANSLATIONS.update(lang_translations)
                                print(f"Successfully loaded translations from {file}")
                        except Exception as e:
                            print(f"Failed to load translation file {file}: {str(e)}")
            except Exception as e:
                print(f"Failed to list translation files: {str(e)}")
        else:
            print(f"Translations directory does not exist and could not be created")
        
        print(f"Available languages after loading: {list(TRANSLATIONS.keys())}")
        return list(TRANSLATIONS.keys())
    except Exception as e:
        print(f"Error in load_translations: {str(e)}")
        return list(TRANSLATIONS.keys())

# Load translations at startup
load_translations()

def initialize_nltk(root):
    """Initialize NLTK data."""
    try:
        # Try to use WordNet to verify it's properly installed
        wordnet.synsets('test')
        return True
    except (LookupError, ImportError, AttributeError) as e:
        print(f"Error initializing NLTK: {str(e)}")
        messagebox.showerror(
            "Error",
            "Failed to initialize word variation detection.\nThis may affect some functionality."
        )
        return False

def get_word_variations(word):
    """Get word variations using NLTK WordNet."""
    if not hasattr(get_word_variations, 'lemmatizer'):
        get_word_variations.lemmatizer = WordNetLemmatizer()
    
    word = word.lower()
    variations = {word}  # Start with the original word
    
    print(f"\nAnalyzing word: {word}")
    
    # Try to find base form (infinitive for verbs)
    base_form = wordnet.morphy(word)
    if base_form:
        variations.add(base_form)
        print(f"Base form: {base_form}")
    else:
        base_form = word
    
    # Get all synsets for the word
    all_synsets = wordnet.synsets(word)
    print(f"Direct synsets: {[s.name() for s in all_synsets]}")
    
    # Process each synset
    for synset in all_synsets:
        print(f"\nProcessing synset: {synset.name()}")
        # Only process lemmas that match our word exactly
        for lemma in synset.lemmas():
            if lemma.name().lower() == word:  # Only process exact matches
                print(f"  Processing variations for: {lemma.name()}")
                variations.add(lemma.name().lower())
                
                # Get derivationally related forms
                deriv_forms = lemma.derivationally_related_forms()
                if deriv_forms:
                    print(f"  Derivational forms: {[d.name() for d in deriv_forms]}")
                    variations.update(d.name().lower() for d in deriv_forms)
    
    # Try to find verb forms if it's a verb
    verb_base = wordnet.morphy(word, 'v')
    if verb_base and verb_base == word:  # Only if the base form matches our word
        # Add common verb forms
        variations.update([
            word + 'ing',  # present participle
            word + 's',    # 3rd person singular
            word + 'ed'    # past tense
        ])
        
        # Handle special cases for verbs ending in 'e'
        if word.endswith('e'):
            variations.add(word[:-1] + 'ing')  # e.g., 'take' -> 'taking'
        
        # Get verb synsets specifically to find irregular forms
        verb_synsets = wordnet.synsets(word, pos='v')
        for synset in verb_synsets:
            for lemma in synset.lemmas():
                if lemma.name().lower() == word:  # Only process exact matches
                    # Try to get irregular forms
                    for form in lemma.derivationally_related_forms():
                        variations.add(form.name().lower())
    
    # Also check common adjective forms if they exist in WordNet
    adj_forms = [word + 'ic', word + 'ical', word + 'ous']
    for adj in adj_forms:
        if wordnet.synsets(adj):
            print(f"Found adjective form: {adj}")
            variations.add(adj)
    
    # Remove any empty strings or None values
    variations = {v for v in variations if v}
    
    print(f"\nFinal variations for {word}: {sorted(variations)}")
    
    return variations

class LexiGen:
    def __init__(self, root):
        self.root = root
        
        # Configure main window
        self.root.title(f"LexiGen v{VERSION} - Fill-in-the-Blank Generator")
        self.root.geometry("1100x800")
        
        # Initialize UI and components
        self.setup_ui()
        
        # Initialize NLTK
        self.nltk_ready = False
        self.init_nltk()
        
        # Check server and fetch models
        self.root.after(100, self.initial_setup)
        
        # Check for updates
        self.root.after(200, lambda: self.check_for_updates(show_message=False))
        
        # Load translations
        available_languages = load_translations()
        if hasattr(self, 'language_select'):
            self.language_select['values'] = available_languages
        
        # Test prompt on startup
        self.root.after(300, self.test_prompt)

    def init_nltk(self):
        """Initialize NLTK after the main UI is ready."""
        self.nltk_ready = initialize_nltk(self.root)
        if self.nltk_ready:
            self.lemmatizer = WordNetLemmatizer()
        else:
            # Retry after a delay if download window was shown
            self.root.after(2000, self.init_nltk)

    def setup_ui(self):
        """Setup the main UI components."""
        # Create styles for update button
        self.style = ttk.Style()
        self.style.configure("Update.TButton", foreground="black")
        self.style.configure("UpdateAvailable.TButton", foreground="blue")
        self.style.configure("UpToDate.TButton", foreground="gray")
        
        # Ollama API Configuration
        self.api_url = "http://127.0.0.1:11434/api/generate"
        self.model = "llama2"
        self.available_models = []
        self.server_connected = False  # Track server connection status
        
        # Default prompt
        self.default_prompt = "Create a simple sentence using the word '{word}'. The sentence should be clear and educational."
        self.prompt_file = "prompt.txt"
        
        # Language selection
        self.current_language = "English"
        
        # Create main container
        self.main_container = ttk.Frame(self.root, padding="10")
        self.main_container.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Settings Frame
        self.settings_frame = ttk.LabelFrame(self.main_container, text=self._("settings"), padding="5")
        self.settings_frame.grid(row=0, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=5)
        
        # Version Label (left-aligned)
        version_label = ttk.Label(self.settings_frame, text=f"v{VERSION}", foreground="gray")
        version_label.grid(row=0, column=0, padx=5, sticky=tk.W)
        
        # Language Selection
        ttk.Label(self.settings_frame, text=self._("language")).grid(row=0, column=1, padx=5)
        self.language_var = tk.StringVar(value=self.current_language)
        self.language_select = ttk.Combobox(self.settings_frame, textvariable=self.language_var, 
                                          values=list(TRANSLATIONS.keys()), state="readonly", width=10)
        self.language_select.grid(row=0, column=2, padx=5)
        self.language_select.bind('<<ComboboxSelected>>', self._on_language_change)
        
        # API URL Entry
        ttk.Label(self.settings_frame, text=self._("api_url")).grid(row=0, column=3, padx=5)
        self.api_url_var = tk.StringVar(value=self.api_url)
        self.api_url_entry = ttk.Entry(self.settings_frame, textvariable=self.api_url_var, width=40)
        self.api_url_entry.grid(row=0, column=4, padx=5)
        
        # Model Selection
        ttk.Label(self.settings_frame, text=self._("model")).grid(row=0, column=5, padx=5)
        self.model_var = tk.StringVar(value=self.model)
        self.model_select = ttk.Combobox(self.settings_frame, textvariable=self.model_var, width=20, state="readonly")
        self.model_select.grid(row=0, column=6, padx=5)
        
        # Server Status and Help
        self.status_frame = ttk.Frame(self.settings_frame)
        self.status_frame.grid(row=1, column=0, columnspan=5, sticky=(tk.W, tk.E), pady=5)
        
        # First row: Status labels
        self.status_labels_frame = ttk.Frame(self.status_frame)
        self.status_labels_frame.pack(fill=tk.X, pady=(0, 5))
        
        self.status_label = ttk.Label(self.status_labels_frame, text="Server Status: Checking...", foreground="gray")
        self.status_label.pack(side=tk.LEFT, padx=5)
        
        # Add separator
        ttk.Separator(self.status_labels_frame, orient='vertical').pack(side=tk.LEFT, padx=5, fill='y')
        
        # Prompt status label
        self.prompt_status_label = ttk.Label(self.status_labels_frame, text=self._("using_default_prompt"), foreground="gray")
        self.prompt_status_label.pack(side=tk.LEFT, padx=5)
        
        # Second row: Buttons
        self.buttons_row = ttk.Frame(self.status_frame)
        self.buttons_row.pack(fill=tk.X)
        
        self.check_server_btn = ttk.Button(self.buttons_row, text="Check Server", command=self.check_server_status)
        self.check_server_btn.pack(side=tk.LEFT, padx=5)
        
        self.help_btn = ttk.Button(self.buttons_row, text="Setup Help", command=self.open_help)
        self.help_btn.pack(side=tk.LEFT, padx=5)
        
        # Create update button but don't pack it yet
        self.update_btn = ttk.Button(self.buttons_row, text="Check for Updates", command=self.check_for_updates)
        
        # Remove the unconditional creation of toggle_prompt_btn
        # We'll create it only if a valid prompt file exists
        self.toggle_prompt_btn = None
        
        # Word Input Frame
        self.input_frame = ttk.LabelFrame(self.main_container, text="Input Words", padding="5")
        self.input_frame.grid(row=1, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=5)
        
        # Word Input
        self.word_input = scrolledtext.ScrolledText(self.input_frame, height=3, width=70)
        self.word_input.grid(row=0, column=0, columnspan=2, padx=5, pady=5)
        self.words_label = ttk.Label(self.input_frame, text=self._("enter_words"))
        self.words_label.grid(row=1, column=0, sticky=tk.W, padx=5)
        
        # Buttons Frame for Generate and Append
        buttons_frame = ttk.Frame(self.input_frame)
        buttons_frame.grid(row=1, column=1, sticky=tk.E, padx=5)
        
        # Generate Button (create first but pack second)
        self.generate_btn = ttk.Button(buttons_frame, text="Generate", command=lambda: self.generate_sentences(append=False))
        
        # Append Button (create second but pack first)
        self.append_btn = ttk.Button(buttons_frame, text="Append", command=lambda: self.generate_sentences(append=True))
        self.append_btn.pack(side=tk.LEFT, padx=(0, 5))  # Pack first to be on the left
        self.append_btn.pack_forget()  # Hide initially
        
        # Now pack the Generate button
        self.generate_btn.pack(side=tk.LEFT)  # Pack second to be on the right
        
        # Progress Bar (hidden by default)
        self.progress_frame = ttk.Frame(self.input_frame)
        self.progress_frame.grid(row=2, column=0, columnspan=2, sticky=(tk.W, tk.E), padx=5, pady=5)
        self.progress_label = ttk.Label(self.progress_frame, text="Generating sentences...")
        self.progress_label.pack(side=tk.LEFT, padx=(0, 5))
        self.progress_bar = ttk.Progressbar(self.progress_frame, mode='determinate')
        self.progress_bar.pack(side=tk.LEFT, fill=tk.X, expand=True)
        self.progress_frame.grid_remove()  # Hide initially
        
        # Sentences Frame
        self.sentences_frame = ttk.LabelFrame(self.main_container, text="Generated Sentences", padding="5")
        self.sentences_frame.grid(row=2, column=0, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), pady=5)
        
        # Buttons Frame
        self.buttons_frame = ttk.Frame(self.sentences_frame)
        self.buttons_frame.grid(row=0, column=0, sticky=tk.E, padx=5, pady=5)

        # Export Docx Button
        self.export_btn = ttk.Button(self.buttons_frame, text="Export Docx", command=self.export_docx, state="disabled")
        self.export_btn.pack(side=tk.LEFT, padx=(0, 5))
        
        # Show All Button
        self.show_all_btn = ttk.Button(self.buttons_frame, text="Show All", command=self.show_all_words, state="disabled")
        self.show_all_btn.pack(side=tk.LEFT, padx=(0, 5))
        
        # Delete All Button
        self.delete_btn = ttk.Button(self.buttons_frame, text="Delete All", command=self.delete_all, state="disabled")
        self.delete_btn.pack(side=tk.LEFT)
        
        # Create a canvas and scrollbar for the sentences
        self.canvas = tk.Canvas(self.sentences_frame)
        self.scrollbar = ttk.Scrollbar(self.sentences_frame, orient="vertical", command=self.canvas.yview)
        self.canvas.configure(yscrollcommand=self.scrollbar.set)
        
        # Container for generated sentences
        self.sentences_container = ttk.Frame(self.canvas)
        self.canvas_frame = self.canvas.create_window((0, 0), window=self.sentences_container, anchor="nw")
        
        # Grid layout for canvas and scrollbar
        self.canvas.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(5, 0))
        self.scrollbar.grid(row=1, column=1, sticky=(tk.N, tk.S))
        
        # List to keep track of generated sentence widgets
        self.sentence_widgets = []
        
        # Configure grid weights
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        self.main_container.columnconfigure(0, weight=1)
        self.main_container.rowconfigure(2, weight=1)  # Make sentences frame expandable
        self.sentences_frame.columnconfigure(0, weight=1)
        self.sentences_frame.rowconfigure(1, weight=1)  # Make canvas expandable
        self.sentences_container.columnconfigure(0, weight=1)
        
        # Bind canvas resizing
        self.sentences_container.bind('<Configure>', self._on_frame_configure)
        self.canvas.bind('<Configure>', self._on_canvas_configure)
        
        # Bind mouse wheel scrolling
        self.canvas.bind_all("<MouseWheel>", self._on_mousewheel)
        
        # Check server status and fetch models on startup
        self.root.after(1000, self.initial_setup)

    def _on_frame_configure(self, event=None):
        """Reset the scroll region to encompass the inner frame"""
        self.canvas.configure(scrollregion=self.canvas.bbox("all"))

    def _on_canvas_configure(self, event):
        """When the canvas is resized, resize the inner frame to match"""
        self.canvas.itemconfig(self.canvas_frame, width=event.width)

    def _on_mousewheel(self, event):
        """Handle mouse wheel scrolling"""
        if self.canvas.winfo_height() < self.sentences_container.winfo_height():
            self.canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

    def check_server_status_initial(self):
        """Check server status on startup, only show message if there's an error."""
        try:
            response = requests.get(self.api_url_var.get().replace("/generate", "/version"))
            if response.status_code == 200:
                self.status_label.config(text=self._("server_status_connected"), foreground="green")
                self.server_connected = True
                self.generate_btn.configure(state="normal")  # Enable Generate button
                return True
            else:
                self.status_label.config(text=self._("server_status_error"), foreground="red")
                messagebox.showerror("Server Status", "Failed to connect to Ollama server.\n\nPlease check if the server is running and the URL is correct.")
                self.server_connected = False
                self.generate_btn.configure(state="disabled")  # Disable Generate button
                return False
        except Exception as e:
            self.status_label.config(text=self._("server_status_not_connected"), foreground="red")
            messagebox.showerror(
                "Server Status", 
                "Cannot connect to Ollama server. Please ensure:\n\n"
                "1. Ollama is installed\n"
                "2. Server is running ('ollama serve')\n"
                "3. API URL is correct\n\n"
                f"Error: {str(e)}"
            )
            self.server_connected = False
            self.generate_btn.configure(state="disabled")  # Disable Generate button
            return False

    def initial_setup(self):
        if self.check_server_status_initial():
            self.fetch_models()

    def fetch_models(self):
        try:
            response = requests.get(self.api_url_var.get().replace("/generate", "/tags"))
            if response.status_code == 200:
                models = response.json()
                self.available_models = [model["name"] for model in models["models"]]
                if not self.available_models:
                    self.available_models = ["llama2"]  # Default fallback
                self.model_select["values"] = self.available_models
                
                # Set default model if current selection is not in list
                if self.model_var.get() not in self.available_models and self.available_models:
                    self.model_var.set(self.available_models[0])
                    
                return True
        except Exception as e:
            self.available_models = ["llama2"]  # Default fallback
            self.model_select["values"] = self.available_models
            self.model_var.set("llama2")
            messagebox.showwarning("Warning", f"Failed to fetch models: {str(e)}\nUsing default model list.")
            return False

    def check_server_status(self):
        try:
            response = requests.get(self.api_url_var.get().replace("/generate", "/version"))
            if response.status_code == 200:
                self.status_label.config(text=self._("server_status_connected"), foreground="green")
                version = response.json().get('version', 'Unknown')
                messagebox.showinfo(
                    self._("server_status_title"),
                    self._("server_connected_msg").format(version=version)
                )
                self.server_connected = True
                self.generate_btn.configure(state="normal")
                return True
            else:
                self.status_label.config(text=self._("server_status_error"), foreground="red")
                messagebox.showerror(
                    self._("server_status_title"),
                    self._("server_error_msg")
                )
                self.server_connected = False
                self.generate_btn.configure(state="disabled")
                return False
        except Exception as e:
            self.status_label.config(text=self._("server_status_not_connected"), foreground="red")
            messagebox.showerror(
                self._("server_status_title"),
                self._("server_connection_error_msg").format(error=str(e))
            )
            self.server_connected = False
            self.generate_btn.configure(state="disabled")
            return False

    def open_help(self):
        """Open the GitHub README page."""
        webbrowser.open("https://github.com/gitmichaelqiu/LexiGen/blob/main/README.md")

    def get_prompt_template(self):
        """Read prompt from file or return default prompt."""
        # If we have a cached prompt, return it directly
        if hasattr(self, '_current_prompt') and self._current_prompt is not None:
            return self._current_prompt
        
        try:
            # Get the application's base directory
            if getattr(sys, 'frozen', False):
                # Running in a bundle (packaged)
                if sys.platform == 'darwin':
                    # For macOS, get the parent directory of the .app bundle
                    base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(sys.executable))))
                else:
                    # For Windows, get the directory containing the .exe
                    base_dir = os.path.dirname(sys.executable)
                print(f"Running in packaged mode, base directory: {base_dir}")
            else:
                # Running in a normal Python environment
                base_dir = os.path.dirname(os.path.abspath(__file__))
                print(f"Running in development mode, base directory: {base_dir}")
            
            prompt_file_path = os.path.join(base_dir, self.prompt_file)
            print(f"Looking for prompt file at: {prompt_file_path}")
            
            if os.path.exists(prompt_file_path):
                try:
                    with open(prompt_file_path, 'r', encoding='utf-8') as f:
                        prompt = f.read().strip()
                        print(f"Read prompt content: {prompt}")
                        
                        # If file exists and is not empty, show Toggle button
                        if prompt:
                            if not self.toggle_prompt_btn:
                                self.toggle_prompt_btn = ttk.Button(
                                    self.buttons_row,
                                    text=self._("toggle_prompt"),
                                    command=self.toggle_prompt
                                )
                                self.toggle_prompt_btn.pack(side=tk.LEFT, padx=5)
                    
                    # Check if prompt contains {word} placeholder
                    if '{word}' in prompt:
                        print(f"Successfully loaded prompt from: {prompt_file_path}")
                        self.prompt_status_label.configure(
                            text=self._("using_custom_prompt"),
                            foreground="green"
                        )
                        self._current_prompt = prompt
                        return prompt
                    else:
                        print(f"Invalid prompt: missing {{word}} placeholder")
                except Exception as e:
                    print(f"Failed to read prompt file: {str(e)}")
            else:
                print(f"Prompt file not found at: {prompt_file_path}")
            
            print("Using default prompt")
            self.prompt_status_label.configure(
                text=self._("using_default_prompt"),
                foreground="gray"
            )
            self._current_prompt = self.default_prompt
            return self.default_prompt
            
        except Exception as e:
            print(f"Error in get_prompt_template: {str(e)}")
            self.prompt_status_label.configure(
                text=self._("using_default_prompt"),
                foreground="gray"
            )
            self._current_prompt = self.default_prompt
            return self.default_prompt

    def check_server_status_silent(self):
        """Check server status without showing message boxes."""
        try:
            response = requests.get(self.api_url_var.get().replace("/generate", "/version"))
            if response.status_code == 200:
                self.status_label.config(text=self._("server_status_connected"), foreground="green")
                self.server_connected = True
                return True
            else:
                self.status_label.config(text=self._("server_status_error"), foreground="red")
                self.server_connected = False
                self.generate_btn.configure(state="disabled")
                return False
        except Exception as e:
            self.status_label.config(text=self._("server_status_not_connected"), foreground="red")
            self.server_connected = False
            self.generate_btn.configure(state="disabled")
            return False

    def generate_sentence_for_word(self, word):
        if not self.check_server_status_silent():
            messagebox.showerror(
                "Server Error",
                "Cannot connect to Ollama server. Please ensure:\n"
                "1. Ollama is installed\n"
                "2. Server is running ('ollama serve')\n"
                "3. Model is installed ('ollama pull llama2')\n"
                "Click 'Setup Help' for more information."
            )
            return None
        
        prompt_template = self.get_prompt_template()
        prompt = prompt_template.format(word=word)
        
        max_attempts = 3
        for attempt in range(max_attempts):
            try:
                response = requests.post(
                    self.api_url_var.get(),
                    json={
                        "model": self.model_var.get(),
                        "prompt": prompt,
                        "stream": False
                    }
                )
                response.raise_for_status()
                result = response.json()
                sentence = result["response"].strip()
                
                # Check if the word or its variations are in the sentence
                word_variations = {word}  # Start with the original word
                # Add simple variations (lowercase, uppercase first letter, all uppercase)
                word_variations.add(word.lower())
                word_variations.add(word.capitalize())
                word_variations.add(word.upper())
                
                # Check if any variation of the word is in the sentence
                if any(variation in sentence.lower() for variation in word_variations):
                    return sentence
                
                # If this is the last attempt, return the sentence anyway
                if attempt == max_attempts - 1:
                    return sentence
                
            except requests.exceptions.RequestException as e:
                messagebox.showerror("Error", f"Failed to generate sentence: {str(e)}\nPlease check your server settings and connection.")
                return None
            except Exception as e:
                messagebox.showerror("Error", f"An unexpected error occurred: {str(e)}")
                return None
        
        return None  # Should never reach here due to return in last attempt

    def create_sentence_widget(self, word, sentence):
        frame = ttk.Frame(self.sentences_container)
        frame.grid(sticky=(tk.W, tk.E), pady=2)
        frame.columnconfigure(0, weight=1)
        
        # Function to mask a word
        def create_mask(w):
            return w[0] + "_" * (len(w) - 1)
        
        # Get word variations (all lowercase)
        word_forms = {w.lower() for w in get_word_variations(word)}
        
        # Create the masked sentence
        sentence_lower = sentence.lower()
        masked_words = {}  # Store original and masked versions
        
        # Sort word forms by length in descending order to handle longer forms first
        word_forms = sorted(word_forms, key=len, reverse=True)
        
        for form in word_forms:
            index = 0
            while True:
                index = sentence_lower.find(form, index)
                if index == -1:
                    break
                    
                # Check if it's a complete word (not part of another word)
                before = index == 0 or not sentence_lower[index-1].isalpha()
                after = index + len(form) == len(sentence_lower) or not sentence_lower[index + len(form)].isalpha()
                
                if before and after:
                    # Get the original case version from the sentence
                    original_word = sentence[index:index + len(form)]
                    masked_word = create_mask(original_word)
                    masked_words[original_word] = masked_word
                
                index += len(form)
        
        # Apply masking while preserving original case
        masked_sentence = sentence
        for original, masked in sorted(masked_words.items(), key=lambda x: len(x[0]), reverse=True):
            masked_sentence = masked_sentence.replace(original, masked)
        
        # Text widget for the sentence
        text_widget = tk.Text(frame, wrap=tk.WORD, cursor="arrow", height=1, width=50)  # Set a reasonable default width
        text_widget.insert("1.0", masked_sentence)
        text_widget.configure(state="disabled", selectbackground=text_widget.cget("background"), 
                            selectforeground=text_widget.cget("foreground"), inactiveselectbackground=text_widget.cget("background"))
        text_widget.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=(0, 5))
        
        # Bind to prevent text selection
        text_widget.bind("<Button-1>", lambda e: "break")
        text_widget.bind("<B1-Motion>", lambda e: "break")
        text_widget.bind("<Double-Button-1>", lambda e: "break")
        text_widget.bind("<Triple-Button-1>", lambda e: "break")
        
        # Disable text widget scrolling
        text_widget.configure(yscrollcommand=None)
        
        # Adjust height based on content
        def adjust_text_height(event=None):
            text_widget.configure(state="normal")
            # Get the number of lines needed to display all text
            text_widget.see("end")  # Make sure all text is rendered
            num_lines = text_widget.count("1.0", "end", "displaylines")[0]
            # Set minimum height to 2 lines
            text_widget.configure(height=max(2, num_lines))
            text_widget.configure(state="disabled")
        
        # Call height adjustment after a short delay to ensure text is properly rendered
        text_widget.after(10, adjust_text_height)
        
        # Bind to Configure event for window resizing
        text_widget.bind('<Configure>', lambda e: text_widget.after(10, adjust_text_height))
        
        # Show/Hide Word Button
        show_btn = ttk.Button(
            frame,
            text=self._("show"),
            command=lambda t=text_widget, b=None: self.toggle_word(word, t, b)
        )
        show_btn.grid(row=0, column=1, padx=(0, 5))
        
        # Copy Button
        copy_btn = ttk.Button(
            frame,
            text=self._("copy"),
            command=lambda t=text_widget: self.copy_sentence(t)
        )
        copy_btn.grid(row=0, column=2, padx=(0, 5))
        
        # Regenerate Button
        regen_btn = ttk.Button(
            frame,
            text="↻",  # Circular arrow symbol for regenerate
            width=2,
            command=lambda w=word, t=text_widget, f=frame: self.regenerate_sentence(w, t, f)
        )
        regen_btn.grid(row=0, column=3, padx=(0, 5))
        
        # Delete Button
        delete_btn = ttk.Button(
            frame,
            text="×",  # Multiplication sign (×) for delete
            width=2,
            command=lambda f=frame: self.delete_sentence(f)
        )
        delete_btn.grid(row=0, column=4)
        
        # Store the original sentence, masked sentence, and button reference for toggling
        text_widget.original_sentence = sentence
        text_widget.masked_sentence = masked_sentence
        text_widget.word_visible = False
        text_widget.show_button = show_btn
        
        # Store text widget reference in the frame for easy access
        frame.text_widget = text_widget
        
        self.sentence_widgets.append(frame)
        return frame

    def copy_sentence(self, text_widget):
        """Copy the current sentence to clipboard."""
        sentence = text_widget.get("1.0", tk.END).strip()
        self.root.clipboard_clear()
        self.root.clipboard_append(sentence)
        self.root.update()  # Required to finalize clipboard update

    def toggle_word(self, word, text_widget, button):
        text_widget.configure(state="normal")
        
        if not text_widget.word_visible:
            # Show the word
            text_widget.delete("1.0", tk.END)
            text_widget.insert("1.0", text_widget.original_sentence)
            text_widget.show_button.configure(text=self._("hide"))
            text_widget.word_visible = True
        else:
            # Hide the word
            text_widget.delete("1.0", tk.END)
            text_widget.insert("1.0", text_widget.masked_sentence)
            text_widget.show_button.configure(text=self._("show"))
            text_widget.word_visible = False
        
        text_widget.configure(state="disabled")

    def export_docx(self):
        """Export sentences to a Word document."""
        if not self.sentence_widgets:
            messagebox.showwarning(self._("warning_title"), self._("no_sentences_warning"))
            return
        
        # Ask user for document title
        title = simpledialog.askstring(
            self._("document_title_prompt"),
            self._("document_title_prompt"),
            initialvalue="Fill in the Blank"
        )
        if not title:  # User cancelled or entered empty string
            return
        
        # Ask user for save location with the title as default filename
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
            title=self._("save_document_title"),
            initialfile=f"{title}.docx"  # Use title as default filename
        )
        
        if not file_path:  # User cancelled
            return
        
        try:
            # Create document
            doc = Document()
            
            # Add title
            title_para = doc.add_paragraph(title)
            title_format = title_para.runs[0].font
            title_format.size = Pt(16)
            title_format.bold = True
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add sentences
            doc.add_paragraph()  # Add space after title
            for i, widget in enumerate(self.sentence_widgets, 1):
                text_widget = widget.text_widget
                sentence = text_widget.masked_sentence
                p = doc.add_paragraph()
                p.add_run(f"{i}. ").bold = True
                p.add_run(sentence)
            
            # Add answer key section
            doc.add_page_break()
            key_title = doc.add_paragraph("Answer Key")
            key_title_format = key_title.runs[0].font
            key_title_format.size = Pt(16)
            key_title_format.bold = True
            key_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()  # Add space after title
            for i, widget in enumerate(self.sentence_widgets, 1):
                text_widget = widget.text_widget
                # Find the masked words by comparing original and masked sentences
                original_words = set()
                for orig, masked in zip(text_widget.original_sentence.split(), text_widget.masked_sentence.split()):
                    if orig != masked and masked.startswith(orig[0]) and '_' in masked:
                        original_words.add(orig)
                
                # Add answer with number
                p = doc.add_paragraph()
                p.add_run(f"{i}. ").bold = True
                p.add_run(", ".join(sorted(original_words)))
            
            # Save document
            doc.save(file_path)
            messagebox.showinfo(self._("export_success_title"), self._("export_success_msg"))
            
        except Exception as e:
            messagebox.showerror(self._("error_title"), str(e))

    def generate_sentences(self, append=False):
        # Get words from input and convert to lowercase
        words = [word.strip().lower() for word in self.word_input.get("1.0", tk.END).strip().split(",")]
        words = [w for w in words if w]  # Remove empty strings
        
        if not words:
            return
        
        # Clear existing sentences if not appending
        if not append:
            self.delete_all()
            
        # Clear input box
        self.word_input.delete("1.0", tk.END)
            
        # Show progress bar
        self.progress_bar['maximum'] = len(words)
        self.progress_bar['value'] = 0
        self.progress_frame.grid()
        self.generate_btn.configure(state="disabled")
        if hasattr(self, 'append_btn'):
            self.append_btn.configure(state="disabled")
        
        # Update UI
        self.root.update()
        
        has_sentences = False
        for i, word in enumerate(words):
            sentence = self.generate_sentence_for_word(word)
            if sentence:
                self.create_sentence_widget(word, sentence)
                has_sentences = True
            else:
                # If sentence generation failed (likely due to server disconnection)
                self.progress_frame.grid_remove()
                return  # Stop generation
            
            # Update progress
            self.progress_bar['value'] = i + 1
            self.progress_label.configure(text=self._("generating") + f" ({i + 1}/{len(words)})")
            self.root.update()
        
        # Enable Show All and Export buttons if sentences were generated
        if has_sentences or len(self.sentence_widgets) > 0:
            self.show_all_btn.configure(text=self._("show_all"), state="normal")
            self.export_btn.configure(state="normal")
            self.delete_btn.configure(state="normal")  # Enable Delete All button
            self.append_btn.pack(side=tk.LEFT, padx=(0, 5))  # Show Append button
        else:
            self.append_btn.pack_forget()  # Hide Append button
            self.delete_btn.configure(state="disabled")  # Disable Delete All button
        
        # Hide progress bar and re-enable generate button only if server is connected
        self.progress_frame.grid_remove()
        if self.server_connected:
            self.generate_btn.configure(state="normal")
            if hasattr(self, 'append_btn'):
                self.append_btn.configure(state="normal")

    def delete_all(self):
        for widget in self.sentence_widgets:
            widget.destroy()
        self.sentence_widgets.clear()
        # Disable all buttons and reset text
        self.show_all_btn.configure(text="Show All", state="disabled")
        self.export_btn.configure(state="disabled")
        self.delete_btn.configure(state="disabled")  # Disable Delete All button
        # Hide Append button
        self.append_btn.pack_forget()

    def show_all_words(self):
        """Show or hide all words in all sentences."""
        if not self.sentence_widgets:
            return
            
        # Determine the action based on any visible word
        show_all = not any(hasattr(widget.text_widget, 'word_visible') and widget.text_widget.word_visible 
                          for widget in self.sentence_widgets)
        
        # Update button text
        self.show_all_btn.configure(text=self._("hide_all") if show_all else self._("show_all"))
        
        # Create a list of valid widgets to process
        valid_widgets = []
        for widget in self.sentence_widgets[:]:  # Create a copy of the list to iterate
            try:
                # Check if widget still exists and has required attributes
                if (widget.winfo_exists() and 
                    hasattr(widget, 'text_widget') and 
                    widget.text_widget.winfo_exists() and
                    hasattr(widget.text_widget, 'original_sentence') and
                    hasattr(widget.text_widget, 'masked_sentence')):
                    valid_widgets.append(widget)
                else:
                    # Remove invalid widget from the list
                    self.sentence_widgets.remove(widget)
            except tk.TclError:
                # Widget has been destroyed, remove it from the list
                self.sentence_widgets.remove(widget)
        
        # Process only valid widgets
        for widget in valid_widgets:
            try:
                text_widget = widget.text_widget
                text_widget.configure(state="normal")
                text_widget.delete("1.0", tk.END)
                
                if show_all:
                    # Show the word
                    text_widget.insert("1.0", text_widget.original_sentence)
                    text_widget.word_visible = True
                    text_widget.show_button.configure(text=self._("hide"))
                else:
                    # Hide the word
                    text_widget.insert("1.0", text_widget.masked_sentence)
                    text_widget.word_visible = False
                    text_widget.show_button.configure(text=self._("show"))
                
                text_widget.configure(state="disabled")
            except tk.TclError:
                # If any error occurs during processing, skip this widget
                continue

    def regenerate_sentence(self, word, text_widget, frame):
        """Generate a new sentence for the given word and update the widget."""
        new_sentence = self.generate_sentence_for_word(word)
        if new_sentence:
            try:
                # Get current position
                current_index = self.sentence_widgets.index(frame)
                
                # Create new frame first
                new_frame = self.create_sentence_widget(word, new_sentence)
                
                # Remove the new frame from the end of sentence_widgets (it was added by create_sentence_widget)
                self.sentence_widgets.pop()
                
                # Remove and destroy old frame
                self.sentence_widgets.pop(current_index)
                frame.destroy()
                
                # Insert new frame at the correct position
                self.sentence_widgets.insert(current_index, new_frame)
                
                # Update grid positions for all widgets
                for i, widget in enumerate(self.sentence_widgets):
                    widget.grid(row=i, column=0, sticky=(tk.W, tk.E), pady=2)
                
                # Update scroll region
                self._on_frame_configure()
                
            except Exception as e:
                print(f"Error during sentence regeneration: {str(e)}")
                # Clean up any remnants of the old frame
                if frame in self.sentence_widgets:
                    self.sentence_widgets.remove(frame)
                try:
                    frame.destroy()
                except:
                    pass
                    
                # Create new widget at the end as fallback
                new_frame = self.create_sentence_widget(word, new_sentence)
                new_frame.grid(row=len(self.sentence_widgets)-1, column=0, sticky=(tk.W, tk.E), pady=2)
                self._on_frame_configure()

    def delete_sentence(self, frame):
        """Delete a single sentence widget."""
        # Remove from the list of sentence widgets
        if frame in self.sentence_widgets:
            self.sentence_widgets.remove(frame)
            
        # Destroy the frame
        frame.destroy()
        
        # Update the scroll region
        self._on_frame_configure()
        
        # If no sentences left, disable buttons and hide append button
        if not self.sentence_widgets:
            self.show_all_btn.configure(text="Show All", state="disabled")
            self.export_btn.configure(state="disabled")
            self.delete_btn.configure(state="disabled")  # Disable Delete All button
            self.append_btn.pack_forget()

    def check_for_updates(self, show_message=True):
        """Check for updates by querying GitHub releases."""
        try:
            response = requests.get("https://api.github.com/repos/gitmichaelqiu/LexiGen/releases/latest")
            if response.status_code == 200:
                latest_version = response.json()["tag_name"].lstrip('v')
                if latest_version > VERSION:
                    self.update_btn.configure(text="New version available", style="UpdateAvailable.TButton")
                else:
                    self.update_btn.configure(text="Up to date", style="UpToDate.TButton")
                
                # Show the update button now that we have a result
                self.update_btn.pack(side=tk.LEFT, padx=5)
                
                if show_message:
                    if latest_version > VERSION:
                        if messagebox.askyesno(
                            "Update Available", 
                            f"A new version (v{latest_version}) is available!\n\nCurrent version: v{VERSION}\n\nWould you like to visit the download page?"
                        ):
                            webbrowser.open("https://github.com/gitmichaelqiu/LexiGen/releases/latest")
                    else:
                        messagebox.showinfo("No Updates", f"You are using the latest version (v{VERSION}).")
            else:
                # Show button with default style if check fails
                self.update_btn.configure(text="Check for Updates", style="Update.TButton")
                self.update_btn.pack(side=tk.LEFT, padx=5)
                if show_message:
                    messagebox.showerror("Error", "Failed to check for updates. Please try again later.")
        except Exception as e:
            # Show button with default style if check fails
            self.update_btn.configure(text="Check for Updates", style="Update.TButton")
            self.update_btn.pack(side=tk.LEFT, padx=5)
            if show_message:
                messagebox.showerror("Error", f"Failed to check for updates:\n{str(e)}")

    def _(self, key):
        """Get translation for the given key."""
        return TRANSLATIONS[self.current_language].get(key, key)

    def _on_language_change(self, event=None):
        """Handle language change event."""
        self.current_language = self.language_var.get()
        self._update_ui_texts()

    def _update_ui_texts(self):
        """Update all UI texts after language change."""
        # Update settings frame
        self.settings_frame.configure(text=self._("settings"))
        
        # Update server status with current state
        if hasattr(self, 'server_connected'):
            if self.server_connected:
                status_text = self._("server_status_connected")
                color = "green"
            else:
                status_text = self._("server_status_not_connected")
                color = "red"
            self.status_label.configure(text=status_text, foreground=color)
        else:
            self.status_label.configure(text=self._("server_status_checking"), foreground="gray")
        
        # Update all labels in settings frame
        for child in self.settings_frame.winfo_children():
            if isinstance(child, ttk.Label):
                text = child.cget("text")
                if text.startswith("v"):  # Version label
                    continue
                elif text.endswith(":") or text.endswith("："):  # Labels ending with colon
                    if "API" in text or "api" in text:
                        child.configure(text=self._("api_url"))
                    elif "Model" in text or "模型" in text:
                        child.configure(text=self._("model"))
                    elif "Language" in text or "语言" in text:
                        child.configure(text=self._("language"))
        
        # Update buttons
        self.check_server_btn.configure(text=self._("check_server"))
        self.help_btn.configure(text=self._("setup_help"))
        
        # Update toggle prompt button if it exists
        if self.toggle_prompt_btn:
            self.toggle_prompt_btn.configure(text=self._("toggle_prompt"))
        
        # Update update button based on its current state
        current_text = self.update_btn.cget("text")
        if current_text in ["New version available", "有新版本"]:
            self.update_btn.configure(text=self._("new_version"))
        elif current_text in ["Up to date", "已是最新"]:
            self.update_btn.configure(text=self._("up_to_date"))
        else:
            self.update_btn.configure(text=self._("check_updates"))
        
        # Update input frame
        self.input_frame.configure(text=self._("input_words"))
        self.generate_btn.configure(text=self._("generate"))
        self.append_btn.configure(text=self._("append"))
        
        # Update sentences frame
        self.sentences_frame.configure(text=self._("generated_sentences"))
        self.export_btn.configure(text=self._("export_docx"))
        self.show_all_btn.configure(text=self._("show_all"))
        self.delete_btn.configure(text=self._("delete_all"))
        
        # Update progress label if visible
        if not self.progress_frame.winfo_ismapped():
            self.progress_label.configure(text=self._("generating"))
        
        # Update the words label
        self.words_label.configure(text=self._("enter_words"))
        
        # Update all sentence widgets
        for frame in self.sentence_widgets:
            for child in frame.winfo_children():
                if isinstance(child, ttk.Button):
                    text = child.cget("text")
                    if text in ["Show", "Hide", "显示", "隐藏"]:
                        child.configure(text=self._("show") if text in ["Show", "显示"] else self._("hide"))
                    elif text in ["Copy", "复制"]:
                        child.configure(text=self._("copy"))
                    elif text == "↻":  # Regenerate button
                        continue
                    elif text == "×":  # Delete button
                        continue
        
        # Update prompt status label
        if hasattr(self, 'prompt_status_label'):
            current_text = self.prompt_status_label.cget("text")
            if current_text == "Using custom prompt":
                self.prompt_status_label.configure(text=self._("using_custom_prompt"))
            else:
                self.prompt_status_label.configure(text=self._("using_default_prompt"))

    def toggle_prompt(self):
        """Toggle between default and custom prompt."""
        current_prompt = self.get_prompt_template()
        if current_prompt == self.default_prompt:
            # Currently using default prompt, try to switch to custom prompt
            prompt_file_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), self.prompt_file)
            if os.path.exists(prompt_file_path):
                with open(prompt_file_path, 'r', encoding='utf-8') as f:
                    prompt = f.read().strip()
                    if '{word}' in prompt:
                        self.prompt_status_label.configure(
                            text=self._("using_custom_prompt"),
                            foreground="green"
                        )
                        # Force get_prompt_template to read file again next time
                        self._current_prompt = None
                    else:
                        messagebox.showwarning(
                            self._("warning_title"),
                            "Invalid prompt format: missing {word} placeholder.\nPlease add {word} to your prompt.txt file."
                        )
            else:
                messagebox.showwarning(
                    self._("warning_title"),
                    "No valid custom prompt found.\nPlease create a prompt.txt file with {word} placeholder."
                )
        else:
            # Currently using custom prompt, switch to default prompt
            self.prompt_status_label.configure(
                text=self._("using_default_prompt"),
                foreground="gray"
            )
            # Force using default prompt
            self._current_prompt = self.default_prompt

    def test_prompt(self):
        """Test prompt loading and display the result."""
        print("\nTesting prompt loading:")
        prompt = self.get_prompt_template()
        print(f"Final prompt: {prompt}")
        if prompt == self.default_prompt:
            print("Using default prompt")
        else:
            print(f"Using custom prompt: {prompt}")

if __name__ == "__main__":
    try:
        root = tk.Tk()
        app = LexiGen(root)
        root.mainloop()
    except Exception as e:
        messagebox.showerror("Error", f"Error starting application:\n{str(e)}")
        import traceback
        traceback.print_exc() 